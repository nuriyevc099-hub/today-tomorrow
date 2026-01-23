from __future__ import annotations
import os
import sys
import shutil
import sqlite3
import json
import unicodedata
from datetime import date, datetime, timedelta
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from flask import Flask, g, render_template, request, redirect, url_for, send_file, flash, session
import random

from sqlalchemy import create_engine

# Word export optional
try:
    from docx import Document
except Exception:
    Document = None

APP_ROOT = Path(__file__).resolve().parent
if getattr(sys, "frozen", False):
    BASE_DIR = Path(sys._MEIPASS)
    DATA_DIR = Path(sys.executable).resolve().parent
else:
    BASE_DIR = APP_ROOT
    DATA_DIR = APP_ROOT

DB_PATH = str(DATA_DIR / "db.sqlite3")
EXPORT_DIR = str(DATA_DIR)
CONSOLE_PIN = "2424"

# (optional) packaged exe üçün: db yoxdursa bundled-dan kopyala
if getattr(sys, "frozen", False):
    bundled_db = BASE_DIR / "db.sqlite3"
    if not Path(DB_PATH).exists() and bundled_db.exists():
        shutil.copy2(bundled_db, DB_PATH)

# =========================
# DB ENGINE (Neon / SQLite)
# =========================
DATABASE_URL = os.environ.get("DATABASE_URL")

if DATABASE_URL:
    # Bəzi platformalarda "postgres://" ola bilir, SQLAlchemy üçün "postgresql://" lazımdır
    if DATABASE_URL.startswith("postgres://"):
        DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

    # Neon üçün SSL parametrləri URL-də artıq olur; pool_pre_ping bağlantını sağlam saxlayır
    engine = create_engine(DATABASE_URL, pool_pre_ping=True)
else:
    # local / fallback
    engine = create_engine(f"sqlite:///{DB_PATH}", pool_pre_ping=True)


# 10 normal slot + 1 bas = 11
GROUP_SLOTS = [
    ("g1", 1),
    ("g2", 1),
    ("g3", 1),
    ("p1", 1),
    ("p2", 1),
    ("p3", 1),
    ("nbm", 2),
    ("nbm2", 2),
]
GROUP_SLOT_COUNTS = {g: c for g, c in GROUP_SLOTS}
NORMAL_SLOTS = sum(GROUP_SLOT_COUNTS.values())
GROUP_PICK_PRIORITY = ["nbm", "p2", "g2", "p3", "g3", "g1", "p1", "nbm2"]
GROUP_PICK_ORDER = [
    g for g in GROUP_PICK_PRIORITY for _ in range(GROUP_SLOT_COUNTS.get(g, 0))
]
GROUP_DIFFICULTY = {
    "nbm": 7,
    "p2": 6,
    "g2": 6,
    "p3": 5,
    "g3": 5,
    "g1": 4,
    "p1": 3,
    "nbm2": 2,
}
MAX_GROUP_DIFFICULTY = max(GROUP_DIFFICULTY.values())
ROTATION_CYCLE = ["g3", "p2", "g2", "nbm2", "p1", "g1", "nbm", "p3"]
ROTATION_INDEX = {g: i for i, g in enumerate(ROTATION_CYCLE)}
ROTATION_BITS = {g: 1 << i for i, g in enumerate(ROTATION_CYCLE)}
ROTATION_FULL_MASK = (1 << len(ROTATION_CYCLE)) - 1

# category map for cooldown logic
GROUP_CATEGORY = {
    "g1": "gun",
    "g2": "gun",
    "g3": "gun",
    "p1": "patrul",
    "p2": "patrul",
    "p3": "patrul",
    "nbm": "nbm",
    "nbm2": "nbm",
    "bas": "bas",
}

# cooldowns (days) to avoid repeating same category/group too soon
CATEGORY_COOLDOWN_DAYS = {"gun": 3, "patrul": 3, "nbm": 4}
GROUP_COOLDOWN_DAYS = {"nbm": 10}  # I NBM should be rare per user
RELAX_PENALTY_CATEGORY = 0.35
RELAX_PENALTY_GROUP = 0.2
MIN_GAP_DAYS = 2  # no back-to-back days for same user
MAX_GAP_DAYS = 21  # cap new/unknown users so they don't dominate
RELAX_PENALTY_YX = 0.35
MONTHLY_GROUP_LIMIT = 3  # max 3 per month per group (normal users)
RELAX_PENALTY_MONTHLY_GROUP = 0.35
BOLUK2_CYCLE_DAYS = 3
TTM_CYCLE_DAYS = 3
NB_COOLDOWN_CHOICES = [5, 6, 7]
ROTATION_GROUPS = [g for g, _ in GROUP_SLOTS]
RELAX_PENALTY_ROTATION = 0.3
RELAX_PENALTY_SAME_CATEGORY = 0.4
RELAX_PENALTY_SAME_GROUP = 0.3
RELAX_PENALTY_NBM_DELAY = 0.45
MAX_YX_PER_DAY = 1
DEFAULT_STATUS_DAYS = {"tm": [], "tp": [], "yx": [5, 6, 7]}
DEFAULT_STATUS_GAP_DAYS = {"ttm": 3}

# statuses:
# aktiv -> normal
# yx -> həftədə 1 dəfə, seçiləndə 5/6/7 gün cooldown (amma yenə də min 7 gün qaydası var)
# izinli/tm/tp/ttm -> seçilməsin
ALWAYS_EXCLUDED_STATUSES = {"izinli", "tm", "tp"}
VALID_ROLES = {"normal", "bas"}
VALID_STATUSES = {"aktiv", "yx", "izinli", "tm", "tp", "ttm"}
REPLACEMENT_EXCLUDED = ALWAYS_EXCLUDED_STATUSES | {"yx"}

STATUS_ALIASES = {"tk": "ttm", "icazeli": "izinli", "icazəli": "izinli"}
AZ_STATUS_TRANSLATION = str.maketrans(
    {
        "\u0131": "i",
        "\u015f": "s",
        "\u011f": "g",
        "\u00e7": "c",
        "\u00f6": "o",
        "\u00fc": "u",
        "\u0259": "e",
    }
)

GROUP_LABELS = {
    "bas": "Tabor Növbətçisi",
    "g1": "I Gün növbətçisi",
    "g2": "II Gün növbətçisi",
    "g3": "III Gün növbətçisi",
    "p1": "I Patrul",
    "p2": "II Patrul",
    "p3": "III Patrul",
    "nbm": "I NBM",
    "nbm2": "II NBM",
}

GROUP_LABELS_SHORT = {
    "bas": "TN",
    "g1": "G1",
    "g2": "G2",
    "g3": "G3",
    "p1": "P1",
    "p2": "P2",
    "p3": "P3",
    "nbm": "NB1",
    "nbm2": "NB2",
}


ROLE_LABELS = {"normal": "Normal", "bas": "Tabor Növbətçisi"}
STATUS_LABELS = {
    "aktiv": "Aktiv",
    "yx": "Yeməkxana",
    "tm": "Tibb Məntəqəsi",
    "tp": "TexnoPark",
    "ttm": "TK",
    "izinli": "İcazəli",
}

SHIFT_START_HOUR = 18  # 18:00
_rng = random.SystemRandom()

app = Flask(
    __name__,
    template_folder=str(BASE_DIR / "templates"),
    static_folder=str(BASE_DIR / "static"),
)
app.secret_key = "local-only-secret"


@app.context_processor
def inject_labels():
    db = get_db()
    settings = load_runtime_settings(db)
    status_days = settings.get("status_days") or {}
    yx_days = status_days.get("yx") or []
    yx_cooldown = settings.get("yx_cooldown_choices") or []
    def fmt_days(days):
        return ", ".join(str(d) for d in days) if days else "-"
    app_info = {
        "boluk2_enabled": settings.get("boluk2_enabled", False),
        "boluk2_cycle_days": settings.get("boluk2_cycle_days", BOLUK2_CYCLE_DAYS),
        "ttm_cycle_days": settings.get("ttm_cycle_days", TTM_CYCLE_DAYS),
        "yx_days": fmt_days(yx_days),
        "yx_cooldown": fmt_days(yx_cooldown),
        "yx_max_per_day": MAX_YX_PER_DAY,
    }
    return dict(
        group_labels=GROUP_LABELS,
        group_labels_short=GROUP_LABELS_SHORT,
        role_labels=ROLE_LABELS,
        status_labels=STATUS_LABELS,
        app_info=app_info,
    )


def _is_postgres() -> bool:
    try:
        return engine.dialect.name == "postgresql"
    except Exception:
        return False


class DBCompat:
    """
    sqlite3 API-yə oxşar wrapper:
    - Postgres-də ? -> %s çevirir
    - INSERT OR IGNORE -> ON CONFLICT DO NOTHING
    - PRAGMA table_info(users) -> Postgres columns query
    - sqlite3.Row tərzi: Postgres-də dict qaytarır (r["col"] işləyir)
    """
    def __init__(self, conn, is_pg: bool):
        self.conn = conn
        self.is_pg = is_pg
        if not is_pg:
            self.conn.row_factory = sqlite3.Row

    def close(self):
        self.conn.close()

    def commit(self):
        self.conn.commit()

    def cursor(self):
        return self.conn.cursor()

    def _rewrite_sql(self, sql: str) -> str:
        if not self.is_pg:
            return sql

        s = sql

        # INSERT OR IGNORE -> ON CONFLICT DO NOTHING
        # (Postgres-də conflict target verməsən də olar)
        s = s.replace("INSERT OR IGNORE INTO", "INSERT INTO")
        if "INSERT INTO" in s and "ON CONFLICT" not in s and "DO NOTHING" not in s:
            # yalnız sqlite-dan gələn "INSERT OR IGNORE" halları üçün
            # (yuxarıda IGNORE silindikdən sonra burada əlavə edirik)
            if "INSERT INTO boluk2_fill_log" in s:
                s = s.strip().rstrip(";") + " ON CONFLICT DO NOTHING"
        return s

    def execute(self, sql: str, params=None):
        # PRAGMA table_info(users) xüsusi hal
        if self.is_pg and sql.strip().upper().startswith("PRAGMA TABLE_INFO("):
            table = sql.strip()[len("PRAGMA table_info("):-1].strip().strip("'").strip('"')
            cur = self.conn.cursor()
            cur.execute(
                """
                SELECT column_name AS name
                FROM information_schema.columns
                WHERE table_schema='public' AND table_name=%s
                ORDER BY ordinal_position
                """,
                (table,),
            )
            rows = cur.fetchall()
            # dict kimi qaytarmaq üçün:
            class _R(dict):
                __getitem__ = dict.get
            return type("C", (), {"fetchall": lambda self2: [_R({"name": r[0]}) for r in rows]})()

        sql2 = self._rewrite_sql(sql)

        if self.is_pg:
            # ? -> %s
            if params is not None:
                sql2 = sql2.replace("?", "%s")

            cur = self.conn.cursor()
            cur.execute(sql2, params or ())
            return cur
        else:
            return self.conn.execute(sql2, params or ())

    def executescript(self, script: str):
        if not self.is_pg:
            cur = self.conn.cursor()
            cur.executescript(script)
            return cur

        # Postgres üçün init_db ayrıca yazılacaq; buranı istifadə etməyəcəyik.
        raise RuntimeError("executescript is not supported for Postgres; use init_db()")


def get_db():
    if "db" not in g:
        raw = engine.raw_connection()
        g.db = DBCompat(raw, _is_postgres())

        if not getattr(g, "status_normalized", False):
            g.status_normalized = True

            # bu PRAGMA çağırışı wrapper ilə Postgres-də də işləyir
            cols = {r["name"] for r in g.db.execute("PRAGMA table_info(users)").fetchall()}

            if "yx_next_eligible" not in cols:
                g.db.execute("ALTER TABLE users ADD COLUMN yx_next_eligible TEXT")
            if "nb_next_eligible" not in cols:
                g.db.execute("ALTER TABLE users ADD COLUMN nb_next_eligible TEXT")
            if "cycle_mask" not in cols:
                g.db.execute("ALTER TABLE users ADD COLUMN cycle_mask INTEGER NOT NULL DEFAULT 0")
            if "cycle_started" not in cols:
                g.db.execute("ALTER TABLE users ADD COLUMN cycle_started TEXT")

            g.db.execute("UPDATE users SET status='ttm' WHERE status='tk'")

            rows = g.db.execute("SELECT id, status FROM users").fetchall()
            for r in rows:
                norm = normalize_status_code(r["status"])
                if norm in VALID_STATUSES and norm != r["status"]:
                    g.db.execute("UPDATE users SET status=? WHERE id=?", (norm, r["id"]))

            # bu CREATE TABLE-lər Postgres-də də işləyəcək (IF NOT EXISTS var)
            g.db.execute(
                """
                CREATE TABLE IF NOT EXISTS boluk2_fill_log (
                    month TEXT NOT NULL,
                    shift_date TEXT NOT NULL,
                    user_id INTEGER NOT NULL,
                    PRIMARY KEY (shift_date, user_id)
                )
                """
            )
            g.db.execute(
                """
                CREATE TABLE IF NOT EXISTS boluk2_fill_history (
                    month TEXT NOT NULL,
                    user_id INTEGER NOT NULL,
                    count INTEGER NOT NULL DEFAULT 0,
                    PRIMARY KEY (month, user_id)
                )
                """
            )
            g.db.execute("CREATE INDEX IF NOT EXISTS idx_boluk2_fill_log_month ON boluk2_fill_log(month)")
            g.db.commit()

    return g.db

@app.teardown_appcontext
def close_db(_exc):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    db = get_db()

    if _is_postgres():
        # Postgres DDL (SQLite AUTOINCREMENT və s. olmadan)
        db.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY,
            name TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'normal',
            status TEXT NOT NULL DEFAULT 'aktiv',
            boluk INTEGER NOT NULL DEFAULT 1,
            last_selected TEXT,
            rotation_score REAL NOT NULL DEFAULT 0,
            last_group TEXT,
            yx_next_eligible TEXT,
            nb_next_eligible TEXT,
            cycle_mask INTEGER NOT NULL DEFAULT 0,
            cycle_started TEXT
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS shifts (
            id BIGSERIAL PRIMARY KEY,
            shift_date TEXT NOT NULL,
            group_name TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            UNIQUE(shift_date, group_name, user_id)
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS shift_changes (
            id BIGSERIAL PRIMARY KEY,
            shift_date TEXT NOT NULL,
            group_name TEXT NOT NULL,
            old_user_id INTEGER NOT NULL,
            new_user_id INTEGER NOT NULL,
            reason TEXT,
            created_at TEXT NOT NULL
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS shift_meta (
            shift_date TEXT PRIMARY KEY,
            confirmed_at TEXT
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS test_meta (
            month TEXT PRIMARY KEY,
            saved_at TEXT
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS test_shifts (
            id BIGSERIAL PRIMARY KEY,
            month TEXT NOT NULL,
            shift_date TEXT NOT NULL,
            group_name TEXT NOT NULL,
            user_id INTEGER NOT NULL
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS weekend_leave (
            user_id INTEGER PRIMARY KEY,
            enabled INTEGER NOT NULL DEFAULT 1
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS boluk2_fill_log (
            month TEXT NOT NULL,
            shift_date TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            PRIMARY KEY (shift_date, user_id)
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS boluk2_fill_history (
            month TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            count INTEGER NOT NULL DEFAULT 0,
            PRIMARY KEY (month, user_id)
        )
        """)

        db.execute("""
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL
        )
        """)

        db.execute("CREATE INDEX IF NOT EXISTS idx_shifts_date ON shifts(shift_date)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_shifts_user ON shifts(user_id)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_shift_meta_date ON shift_meta(shift_date)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_test_shifts_month ON test_shifts(month)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_users_role ON users(role)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_shift_changes_date ON shift_changes(shift_date)")
        db.execute("CREATE INDEX IF NOT EXISTS idx_boluk2_fill_log_month ON boluk2_fill_log(month)")

        db.commit()

        # Seed 40 users if empty
        c = db.execute("SELECT COUNT(*) AS c FROM users").fetchone()
        if int(c["c"]) == 0:
            for i in range(1, 41):
                db.execute(
                    "INSERT INTO users (id, name, role, status, boluk) VALUES (?, ?, 'normal', 'aktiv', 1) ON CONFLICT (id) DO NOTHING",
                    (i, f"User {i}"),
                )
            db.commit()

        return

    # --- SQLite yolu (sənin köhnə məntiqin) ---
    raw = engine.raw_connection()
    raw.row_factory = sqlite3.Row
    cur = raw.cursor()
    cur.executescript(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY,
            name TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'normal',
            status TEXT NOT NULL DEFAULT 'aktiv',
            boluk INTEGER NOT NULL DEFAULT 1,
            last_selected TEXT,
            rotation_score REAL NOT NULL DEFAULT 0,
            last_group TEXT,
            yx_next_eligible TEXT,
            nb_next_eligible TEXT,
            cycle_mask INTEGER NOT NULL DEFAULT 0,
            cycle_started TEXT
        );

        CREATE TABLE IF NOT EXISTS shifts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            shift_date TEXT NOT NULL,
            group_name TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            UNIQUE(shift_date, group_name, user_id)
        );

        CREATE TABLE IF NOT EXISTS shift_changes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            shift_date TEXT NOT NULL,
            group_name TEXT NOT NULL,
            old_user_id INTEGER NOT NULL,
            new_user_id INTEGER NOT NULL,
            reason TEXT,
            created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS shift_meta (
            shift_date TEXT PRIMARY KEY,
            confirmed_at TEXT
        );

        CREATE TABLE IF NOT EXISTS test_meta (
            month TEXT PRIMARY KEY,
            saved_at TEXT
        );

        CREATE TABLE IF NOT EXISTS test_shifts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            month TEXT NOT NULL,
            shift_date TEXT NOT NULL,
            group_name TEXT NOT NULL,
            user_id INTEGER NOT NULL
        );

        CREATE TABLE IF NOT EXISTS weekend_leave (
            user_id INTEGER PRIMARY KEY,
            enabled INTEGER NOT NULL DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS boluk2_fill_log (
            month TEXT NOT NULL,
            shift_date TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            PRIMARY KEY (shift_date, user_id)
        );

        CREATE TABLE IF NOT EXISTS boluk2_fill_history (
            month TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            count INTEGER NOT NULL DEFAULT 0,
            PRIMARY KEY (month, user_id)
        );

        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT NOT NULL
        );

        CREATE INDEX IF NOT EXISTS idx_shifts_date ON shifts(shift_date);
        CREATE INDEX IF NOT EXISTS idx_shifts_user ON shifts(user_id);
        CREATE INDEX IF NOT EXISTS idx_shift_meta_date ON shift_meta(shift_date);
        CREATE INDEX IF NOT EXISTS idx_test_shifts_month ON test_shifts(month);
        CREATE INDEX IF NOT EXISTS idx_users_role ON users(role);
        CREATE INDEX IF NOT EXISTS idx_shift_changes_date ON shift_changes(shift_date);
        CREATE INDEX IF NOT EXISTS idx_boluk2_fill_log_month ON boluk2_fill_log(month);
        """
    )
    raw.commit()
    raw.close()



def get_setting(db: sqlite3.Connection, key: str, default=None):
    row = db.execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
    if not row:
        return default
    try:
        return json.loads(row["value"])
    except Exception:
        return default


def set_setting(db: sqlite3.Connection, key: str, value) -> None:
    payload = json.dumps(value, ensure_ascii=False)
    db.execute(
        """
        INSERT INTO settings (key, value)
        VALUES (?, ?)
        ON CONFLICT(key) DO UPDATE SET value=excluded.value
        """,
        (key, payload),
    )
    db.commit()


def get_all_settings(db: sqlite3.Connection) -> Dict[str, object]:
    rows = db.execute("SELECT key, value FROM settings ORDER BY key ASC").fetchall()
    out: Dict[str, object] = {}
    for r in rows:
        try:
            out[r["key"]] = json.loads(r["value"])
        except Exception:
            out[r["key"]] = r["value"]
    return out


def coerce_int(value, default: int, min_val: int, max_val: int) -> int:
    try:
        n = int(value)
    except Exception:
        return default
    if n < min_val or n > max_val:
        return default
    return n


def normalize_int_list(value, min_val: int, max_val: int, default: List[int]) -> List[int]:
    items: List[int] = []
    if isinstance(value, str):
        parts = [p.strip() for p in value.replace(",", " ").split() if p.strip()]
        for p in parts:
            try:
                items.append(int(p))
            except Exception:
                continue
    elif isinstance(value, (list, tuple)):
        for p in value:
            try:
                items.append(int(p))
            except Exception:
                continue
    if not items:
        return default
    out = sorted({n for n in items if min_val <= n <= max_val})
    return out or default


def normalize_int_map(value, min_val: int, max_val: int) -> Dict[str, int]:
    if not isinstance(value, dict):
        return {}
    out: Dict[str, int] = {}
    for k, v in value.items():
        try:
            n = int(v)
        except Exception:
            continue
        if min_val <= n <= max_val:
            out[str(k)] = n
    return out


def normalize_status_days(value) -> Dict[str, List[int]]:
    if not isinstance(value, dict):
        return {}
    out: Dict[str, List[int]] = {}
    for k, v in value.items():
        key = normalize_status_code(str(k))
        if v is None:
            out[key] = []
            continue
        if isinstance(v, str) and v.strip().lower() in {"none", "off", "never", "0"}:
            out[key] = []
            continue
        if isinstance(v, (list, tuple)) and len(v) == 0:
            out[key] = []
            continue
        days = normalize_int_list(v, 1, 7, [])
        out[key] = days
    return out


def normalize_status_gap(value) -> Dict[str, int]:
    if not isinstance(value, dict):
        return {}
    out: Dict[str, int] = {}
    for k, v in value.items():
        key = normalize_status_code(str(k))
        if key not in VALID_STATUSES:
            continue
        try:
            n = int(v)
        except Exception:
            continue
        if 0 <= n <= 30:
            out[key] = n
    return out


def merge_int_map(defaults: Dict[str, int], override: Dict[str, int]) -> Dict[str, int]:
    out = dict(defaults)
    out.update(override)
    return out


def normalize_status_code(value: str) -> str:
    if value is None:
        return ""
    s = str(value).strip().lower()
    if not s:
        return ""
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = s.translate(AZ_STATUS_TRANSLATION)
    return STATUS_ALIASES.get(s, s)


def load_runtime_settings(db: sqlite3.Connection) -> Dict[str, object]:
    status_days = normalize_status_days(get_setting(db, "status_days", DEFAULT_STATUS_DAYS))
    status_days["tm"] = []
    status_days["tp"] = []
    status_days["yx"] = [5, 6, 7]
    status_days.pop("ttm", None)
    status_gap_days = merge_int_map(
        DEFAULT_STATUS_GAP_DAYS,
        normalize_status_gap(get_setting(db, "status_gap_days", {})),
    )
    min_gap_days = coerce_int(get_setting(db, "min_gap_days", MIN_GAP_DAYS), MIN_GAP_DAYS, 0, 30)
    monthly_group_limit = coerce_int(
        get_setting(db, "monthly_group_limit", MONTHLY_GROUP_LIMIT),
        MONTHLY_GROUP_LIMIT,
        0,
        31,
    )
    cat_override = normalize_int_map(
        get_setting(db, "category_cooldown_days", {}), 0, 30
    )
    group_override = normalize_int_map(
        get_setting(db, "group_cooldown_days", {}), 0, 60
    )
    category_cooldown_days = merge_int_map(CATEGORY_COOLDOWN_DAYS, cat_override)
    group_cooldown_days = merge_int_map(GROUP_COOLDOWN_DAYS, group_override)
    yx_cooldown_choices = normalize_int_list(
        get_setting(db, "yx_cooldown_choices", [5, 6, 7]), 1, 30, [5, 6, 7]
    )
    nb_cooldown_choices = normalize_int_list(
        get_setting(db, "nb_cooldown_choices", NB_COOLDOWN_CHOICES),
        1,
        30,
        NB_COOLDOWN_CHOICES,
    )
    boluk2_enabled = bool(get_setting(db, "boluk2_enabled", False))
    boluk2_cycle_days = coerce_int(
        get_setting(db, "boluk2_cycle_days", BOLUK2_CYCLE_DAYS),
        BOLUK2_CYCLE_DAYS,
        1,
        30,
    )
    boluk2_cycle_days = max(boluk2_cycle_days, BOLUK2_CYCLE_DAYS)
    ttm_cycle_days = coerce_int(
        get_setting(db, "ttm_cycle_days", status_gap_days.get("ttm", TTM_CYCLE_DAYS)),
        TTM_CYCLE_DAYS,
        1,
        30,
    )
    ttm_cycle_days = max(ttm_cycle_days, TTM_CYCLE_DAYS)
    status_gap_days["ttm"] = ttm_cycle_days
    return {
        "status_days": status_days,
        "status_gap_days": status_gap_days,
        "min_gap_days": min_gap_days,
        "monthly_group_limit": monthly_group_limit,
        "category_cooldown_days": category_cooldown_days,
        "group_cooldown_days": group_cooldown_days,
        "yx_cooldown_choices": yx_cooldown_choices,
        "nb_cooldown_choices": nb_cooldown_choices,
        "boluk2_enabled": boluk2_enabled,
        "boluk2_cycle_days": boluk2_cycle_days,
        "ttm_cycle_days": ttm_cycle_days,
    }


def get_boluk2_last_cycle(db: sqlite3.Connection) -> Optional[date]:
    raw = get_setting(db, "boluk2_last_cycle", None)
    if not raw:
        return None
    return parse_date(raw)


def weekday_number(d: date) -> int:
    return d.weekday() + 1


def is_weekend(d: date) -> bool:
    return d.weekday() >= 5


def get_weekend_leave_ids(db: sqlite3.Connection) -> set:
    rows = db.execute(
        "SELECT user_id FROM weekend_leave WHERE enabled=1"
    ).fetchall()
    return {r["user_id"] for r in rows}


def weekend_leave_ok(user_id: int, shift_date: date, weekend_leave_ids: set) -> bool:
    if not weekend_leave_ids:
        return True
    if not is_weekend(shift_date):
        return True
    return user_id not in weekend_leave_ids


def status_allowed_today(status: str, shift_date: date, status_days: Dict[str, List[int]]) -> bool:
    if not status_days:
        return True
    status = normalize_status_code(status)
    if status not in status_days:
        return True
    days = status_days.get(status)
    if days is None:
        return True
    if len(days) == 0:
        return False
    return weekday_number(shift_date) in days


def is_boluk2_cycle_day(
    shift_date: date, last_cycle: Optional[date], cycle_days: int = BOLUK2_CYCLE_DAYS
) -> bool:
    cycle_days = coerce_int(cycle_days, BOLUK2_CYCLE_DAYS, 1, 30)
    if last_cycle is None:
        return True
    if shift_date <= last_cycle:
        return False
    return (shift_date - last_cycle).days >= cycle_days


def get_boluk2_last_cycle_before(
    db: sqlite3.Connection, shift_date: date
) -> Optional[date]:
    boluk2_active = db.execute(
        """
        SELECT id FROM users
        WHERE role='normal' AND boluk=2 AND LOWER(TRIM(status))='aktiv'
        """
    ).fetchall()
    boluk2_active_ids = {r["id"] for r in boluk2_active}
    if not boluk2_active_ids:
        return None
    ds = shift_date.strftime("%Y-%m-%d")
    row = db.execute(
        """
        SELECT s.shift_date AS shift_date
        FROM shifts s
        JOIN users u ON u.id = s.user_id
        WHERE s.shift_date <= ?
          AND u.role='normal'
          AND u.boluk=2
          AND LOWER(TRIM(u.status))='aktiv'
        GROUP BY s.shift_date
        HAVING COUNT(DISTINCT s.user_id) >= ?
        ORDER BY s.shift_date DESC
        LIMIT 1
        """,
        (ds, len(boluk2_active_ids)),
    ).fetchone()
    return parse_date(row["shift_date"]) if row else None


def get_ttm_last_cycle_before(
    db: sqlite3.Connection, shift_date: date
) -> Optional[date]:
    ttm_active = db.execute(
        """
        SELECT id FROM users
        WHERE role='normal' AND LOWER(TRIM(status)) IN ('ttm', 'tk')
        """
    ).fetchall()
    ttm_active_ids = {r["id"] for r in ttm_active}
    if not ttm_active_ids:
        return None
    ds = shift_date.strftime("%Y-%m-%d")
    row = db.execute(
        """
        SELECT s.shift_date AS shift_date
        FROM shifts s
        JOIN users u ON u.id = s.user_id
        WHERE s.shift_date <= ?
          AND u.role='normal'
          AND LOWER(TRIM(u.status)) IN ('ttm', 'tk')
        GROUP BY s.shift_date
        HAVING COUNT(DISTINCT s.user_id) >= ?
        ORDER BY s.shift_date DESC
        LIMIT 1
        """,
        (ds, len(ttm_active_ids)),
    ).fetchone()
    return parse_date(row["shift_date"]) if row else None


def get_prev_cycle_boluk1_ids(db: sqlite3.Connection, cycle_date: date) -> set:
    ds = cycle_date.strftime("%Y-%m-%d")
    rows = db.execute(
        """
        SELECT s.user_id
        FROM shifts s
        JOIN users u ON u.id = s.user_id
        WHERE s.shift_date=? AND u.role='normal' AND u.boluk=1
        """,
        (ds,),
    ).fetchall()
    return {r["user_id"] for r in rows}


def swap_in_forced_users(
    picks: Dict[str, List[int]],
    forced_ids: set,
    users_by_id: Dict[int, sqlite3.Row],
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]] = None,
) -> None:
    assigned = {uid for ids in picks.values() for uid in ids}
    remaining = [uid for uid in forced_ids if uid not in assigned]
    if not remaining:
        return

    def try_swap(
        uid: int,
        require_rotation: bool,
        prefer_boluk1: bool,
        enforce_missing: bool,
        enforce_cycle: bool,
    ) -> bool:
        u = users_by_id.get(uid)
        if not u:
            return False
        missing = user_missing_groups(monthly_group_counts, u["id"])
        for group in GROUP_PICK_ORDER:
            if group not in picks:
                continue
            if enforce_missing and missing and group not in missing:
                continue
            if enforce_cycle and cycle_bit(group):
                if user_cycle_mask(u) & cycle_bit(group):
                    continue
            if require_rotation and not rotation_step_ok(u["last_group"], group):
                continue
            ids = picks[group]
            for i, other_id in enumerate(ids):
                if other_id in forced_ids:
                    continue
                other = users_by_id.get(other_id)
                if not other:
                    continue
                if prefer_boluk1 and user_boluk(other) != 1:
                    continue
                ids[i] = uid
                return True
        return False

    for uid in remaining:
        if try_swap(uid, True, True, True, True):
            continue
        if try_swap(uid, True, False, True, True):
            continue
        if try_swap(uid, False, True, True, True):
            continue
        if try_swap(uid, False, False, True, True):
            continue
        if try_swap(uid, True, True, False, True):
            continue
        if try_swap(uid, True, False, False, True):
            continue
        if try_swap(uid, False, True, False, True):
            continue
        if try_swap(uid, False, False, False, True):
            continue
        if try_swap(uid, True, True, True, False):
            continue
        if try_swap(uid, True, False, True, False):
            continue
        if try_swap(uid, False, True, True, False):
            continue
        if try_swap(uid, False, False, True, False):
            continue
        if try_swap(uid, True, True, False, False):
            continue
        if try_swap(uid, True, False, False, False):
            continue
        if try_swap(uid, False, True, False, False):
            continue
        try_swap(uid, False, False, False, False)


def parse_int_list_arg(raw: str, min_val: int, max_val: int) -> List[int]:
    parts = [p.strip() for p in raw.replace(",", " ").split() if p.strip()]
    if not parts:
        raise ValueError("Missing numbers.")
    out = []
    for p in parts:
        n = int(p)
        if n < min_val or n > max_val:
            raise ValueError(f"Out of range: {n}")
        out.append(n)
    return sorted(set(out))


def command_help_text() -> str:
    return (
        "Commands:\n"
        "  help\n"
        "  show settings\n"
        "  show effective\n"
        "  reset settings\n"
        "  set status_days yx=5,6,7\n"
        "  set status_gap ttm=3\n"
        "  set boluk2_cycle 3\n"
        "\n"
        "Notes:\n"
        "  - Weekdays are 1..7 (1=Mon, 7=Sun).\n"
        "  - izinli is always excluded from selection.\n"
    )


def run_console_command(line: str, db: sqlite3.Connection) -> str:
    tokens = [t for t in line.strip().split() if t.strip()]
    if not tokens:
        return ""
    cmd = tokens[0].lower()
    args = tokens[1:]

    if cmd in {"help", "?"}:
        return command_help_text()

    if cmd == "show" and len(args) == 1 and args[0].lower() == "settings":
        return json.dumps(get_all_settings(db), indent=2, ensure_ascii=False)

    if cmd == "show" and len(args) == 1 and args[0].lower() == "effective":
        return json.dumps(load_runtime_settings(db), indent=2, ensure_ascii=False)

    if cmd == "reset" and len(args) == 1 and args[0].lower() == "settings":
        db.execute("DELETE FROM settings")
        db.commit()
        return "OK: settings cleared."

    if cmd == "set" and len(args) >= 2:
        key = args[0].lower()
        rest = args[1:]
        locked = {
            "min_gap",
            "category_cooldown",
            "group_cooldown",
            "monthly_group_limit",
            "yx_cooldown",
        }
        if key in locked:
            raise ValueError("This setting is locked.")

        if key == "min_gap":
            if len(rest) != 1:
                raise ValueError("Usage: set min_gap 2")
            n = int(rest[0])
            if n < 0 or n > 30:
                raise ValueError("min_gap must be 0..30")
            set_setting(db, "min_gap_days", n)
            return f"OK: min_gap_days={n}"

        if key == "monthly_group_limit":
            if len(rest) != 1:
                raise ValueError("Usage: set monthly_group_limit 3")
            n = int(rest[0])
            if n < 0 or n > 31:
                raise ValueError("monthly_group_limit must be 0..31")
            set_setting(db, "monthly_group_limit", n)
            return f"OK: monthly_group_limit={n}"

        if key == "yx_cooldown":
            if len(rest) != 1:
                raise ValueError("Usage: set yx_cooldown 5,6,7")
            days = parse_int_list_arg(rest[0], 1, 30)
            set_setting(db, "yx_cooldown_choices", days)
            return f"OK: yx_cooldown_choices={days}"

        if key == "status_days":
            if len(rest) == 1 and rest[0].lower() in {"clear", "none"}:
                set_setting(db, "status_days", {})
                return "OK: status_days cleared."
            status_map: Dict[str, List[int]] = {}
            for item in rest:
                if "=" not in item:
                    raise ValueError("status_days expects status=5,6,7")
                s_key, raw = item.split("=", 1)
                s_key = normalize_status_code(s_key.strip().lower())
                raw = raw.strip().lower()
                if raw in {"none", "off", "never", "0"}:
                    days = []
                else:
                    days = parse_int_list_arg(raw, 1, 7)
                if s_key in {"all", "*"}:
                    for st in VALID_STATUSES:
                        status_map[st] = days
                    continue
                if s_key not in VALID_STATUSES:
                    raise ValueError(f"Unknown status: {s_key}")
                status_map[s_key] = days
            set_setting(db, "status_days", status_map)
            return f"OK: status_days={status_map}"

        if key == "status_gap":
            if len(rest) == 1 and rest[0].lower() in {"clear", "none"}:
                set_setting(db, "status_gap_days", {})
                return "OK: status_gap_days cleared."
            status_map: Dict[str, int] = {}
            for item in rest:
                if "=" not in item:
                    raise ValueError("status_gap expects status=days")
                s_key, raw = item.split("=", 1)
                s_key = normalize_status_code(s_key.strip().lower())
                raw = raw.strip().lower()
                if raw in {"none", "off", "0"}:
                    n = 0
                else:
                    n = int(raw)
                if n < 0 or n > 30:
                    raise ValueError("status_gap days must be 0..30")
                if s_key in {"all", "*"}:
                    for st in VALID_STATUSES:
                        if n > 0:
                            status_map[st] = n
                        else:
                            status_map.pop(st, None)
                    continue
                if s_key not in VALID_STATUSES:
                    raise ValueError(f"Unknown status: {s_key}")
                if n > 0:
                    status_map[s_key] = n
            set_setting(db, "status_gap_days", status_map)
            return f"OK: status_gap_days={status_map}"

        if key in {"boluk2_cycle", "boluk2_cycle_days"}:
            if len(rest) != 1:
                raise ValueError("Usage: set boluk2_cycle 2")
            n = int(rest[0])
            if n < 1 or n > 30:
                raise ValueError("boluk2_cycle must be 1..30")
            set_setting(db, "boluk2_cycle_days", n)
            return f"OK: boluk2_cycle_days={n}"

        if key == "category_cooldown":
            if not rest:
                raise ValueError("Usage: set category_cooldown gun=3 patrul=3 nbm=4")
            valid = {"gun", "patrul", "nbm"}
            out: Dict[str, int] = {}
            for item in rest:
                if "=" not in item:
                    raise ValueError("category_cooldown expects key=value")
                c_key, raw = item.split("=", 1)
                c_key = c_key.strip().lower()
                if c_key not in valid:
                    raise ValueError(f"Unknown category: {c_key}")
                n = int(raw)
                if n < 0 or n > 30:
                    raise ValueError("category_cooldown must be 0..30")
                out[c_key] = n
            set_setting(db, "category_cooldown_days", out)
            return f"OK: category_cooldown_days={out}"

        if key == "group_cooldown":
            if not rest:
                raise ValueError("Usage: set group_cooldown nbm=10 nbm2=8")
            valid = {g for g, _ in GROUP_SLOTS} | {"bas"}
            out: Dict[str, int] = {}
            for item in rest:
                if "=" not in item:
                    raise ValueError("group_cooldown expects key=value")
                g_key, raw = item.split("=", 1)
                g_key = g_key.strip().lower()
                if g_key not in valid:
                    raise ValueError(f"Unknown group: {g_key}")
                n = int(raw)
                if n < 0 or n > 60:
                    raise ValueError("group_cooldown must be 0..60")
                out[g_key] = n
            set_setting(db, "group_cooldown_days", out)
            return f"OK: group_cooldown_days={out}"

        raise ValueError("Unknown setting key.")

    raise ValueError("Unknown command. Type 'help'.")


def execute_console_commands(text: str, db: sqlite3.Connection) -> str:
    outputs: List[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        try:
            result = run_console_command(line, db)
        except Exception as exc:
            result = f"ERROR: {exc}"
        outputs.append(f"$ {line}\n{result}")
    return "\n\n".join(outputs)


def parse_date(s: Optional[str]) -> Optional[date]:
    if not s:
        return None
    return datetime.strptime(s, "%Y-%m-%d").date()


def days_between(a: Optional[date], b: date) -> int:
    """b - a in 'shift days' (date diff). If a None -> huge."""
    if a is None:
        return MAX_GAP_DAYS
    return min((b - a).days, MAX_GAP_DAYS)

def already_assigned(user_id: int, picks: Dict[str, List[int]]) -> bool:
    return any(user_id in ids for ids in picks.values())


def month_range(ym: str) -> Tuple[str, str]:
    start = datetime.strptime(ym + "-01", "%Y-%m-%d").date()
    if start.month == 12:
        nxt = date(start.year + 1, 1, 1)
    else:
        nxt = date(start.year, start.month + 1, 1)
    return start.strftime("%Y-%m-%d"), nxt.strftime("%Y-%m-%d")


def previous_month(ym: str) -> str:
    start = datetime.strptime(ym + "-01", "%Y-%m-%d").date()
    if start.month == 1:
        prev = date(start.year - 1, 12, 1)
    else:
        prev = date(start.year, start.month - 1, 1)
    return prev.strftime("%Y-%m")


def get_monthly_counts(ym: str) -> Dict[int, int]:
    db = get_db()
    start, end = month_range(ym)
    rows = db.execute(
        """
        SELECT user_id, COUNT(*) AS c
        FROM shifts
        WHERE shift_date >= ? AND shift_date < ?
        GROUP BY user_id
        """,
        (start, end),
    ).fetchall()
    return {r["user_id"]: r["c"] for r in rows}


def get_monthly_group_counts(ym: str) -> Dict[int, Dict[str, int]]:
    db = get_db()
    start, end = month_range(ym)
    rows = db.execute(
        """
        SELECT user_id, group_name, COUNT(*) AS c
        FROM shifts
        WHERE shift_date >= ? AND shift_date < ?
        GROUP BY user_id, group_name
        """,
        (start, end),
    ).fetchall()
    out: Dict[int, Dict[str, int]] = {}
    for r in rows:
        out.setdefault(r["user_id"], {})[r["group_name"]] = r["c"]
    return out


def monthly_group_count(
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]],
    user_id: int,
    group_name: str,
) -> int:
    if not monthly_group_counts:
        return 0
    return monthly_group_counts.get(user_id, {}).get(group_name, 0)


def any_group_overflow(
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]], user_id: int
) -> bool:
    if not monthly_group_counts:
        return False
    counts = monthly_group_counts.get(user_id, {})
    return any(counts.get(g, 0) > 0 for g in ROTATION_GROUPS)


def monthly_group_factor(count: int) -> float:
    if count <= 0:
        return 1.25
    if count == 1:
        return 1.0
    if count == 2:
        return 0.8
    return 0.6


def user_missing_groups(
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]], user_id: int
) -> List[str]:
    if not monthly_group_counts:
        return list(ROTATION_GROUPS)
    counts = monthly_group_counts.get(user_id, {})
    return [g for g in ROTATION_GROUPS if counts.get(g, 0) == 0]


def monthly_category_counts(
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]], user_id: int
) -> Dict[str, int]:
    counts = {"gun": 0, "patrul": 0, "nbm": 0}
    if not monthly_group_counts:
        return counts
    for group, cnt in monthly_group_counts.get(user_id, {}).items():
        cat = group_category(group)
        if cat in counts:
            counts[cat] += cnt
    return counts


def monthly_difficulty_score(
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]], user_id: int
) -> float:
    if not monthly_group_counts:
        return 0.0
    total = 0.0
    for group, cnt in monthly_group_counts.get(user_id, {}).items():
        total += GROUP_DIFFICULTY.get(group, 0) * cnt
    return total


def difficulty_balance_factor(
    target_group: str,
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]],
    user_id: int,
) -> float:
    if not monthly_group_counts:
        return 1.0
    weight = GROUP_DIFFICULTY.get(target_group, 0)
    if weight <= 0:
        return 1.0
    score = monthly_difficulty_score(monthly_group_counts, user_id)
    scale = weight / float(MAX_GROUP_DIFFICULTY)
    return max(0.35, 1.0 - (score * 0.03 * scale))


def transition_multiplier(user: sqlite3.Row, target_group: str) -> float:
    last_group = user["last_group"]
    if not last_group:
        return 1.0
    last_cat = group_category(last_group)
    target_cat = group_category(target_group)
    if last_cat == "gun":
        if target_cat == "patrul":
            return 1.15
        if target_cat == "nbm":
            return 0.85
        if target_cat == "gun":
            return 0.2
    return 1.0


def rotation_cycle_multiplier(user: sqlite3.Row, target_group: str) -> float:
    last_group = user["last_group"]
    if not last_group:
        return 1.0
    if last_group not in ROTATION_INDEX or target_group not in ROTATION_INDEX:
        return 1.0
    cycle_len = len(ROTATION_CYCLE)
    dist = (ROTATION_INDEX[target_group] - ROTATION_INDEX[last_group]) % cycle_len
    if dist == 1:
        return 1.6
    if dist == 2:
        return 1.1
    if dist == 3:
        return 0.9
    if dist >= cycle_len - 1:
        return 0.6
    return 0.8

def get_rolling_group_counts(
    db: sqlite3.Connection, shift_date: date, window_days: int = 30
) -> Dict[int, Dict[str, int]]:
    start = (shift_date - timedelta(days=window_days)).strftime("%Y-%m-%d")
    end = shift_date.strftime("%Y-%m-%d")
    rows = db.execute(
        """
        SELECT user_id, group_name, COUNT(*) AS c
        FROM shifts
        WHERE shift_date >= ? AND shift_date < ?
        GROUP BY user_id, group_name
        """,
        (start, end),
    ).fetchall()
    counts: Dict[int, Dict[str, int]] = {}
    for r in rows:
        gname = _norm_rot_group(r["group_name"])
        if not gname:
            continue
        counts.setdefault(r["user_id"], {}).setdefault(gname, 0)
        counts[r["user_id"]][gname] += int(r["c"])
    return counts

def build_sim_rolling_group_counts(
    users: List[dict], days: List[date], matrix: Dict[int, Dict[str, str]], day_index: int
) -> Dict[int, Dict[str, int]]:
    start = max(0, day_index - 30)
    counts: Dict[int, Dict[str, int]] = {}
    for i in range(start, day_index):
        ds = days[i].strftime("%Y-%m-%d")
        for u in users:
            g = matrix.get(u["id"], {}).get(ds)
            if not g:
                continue
            gname = _norm_rot_group(g)
            if not gname:
                continue
            counts.setdefault(u["id"], {}).setdefault(gname, 0)
            counts[u["id"]][gname] += 1
    return counts

def compute_rolling_group_avg(
    rolling_counts: Dict[int, Dict[str, int]], eligible_ids: set
) -> Dict[str, float]:
    totals: Dict[str, int] = {}
    for uid in eligible_ids:
        for g, c in rolling_counts.get(uid, {}).items():
            totals[g] = totals.get(g, 0) + int(c)
    denom = max(1, len(eligible_ids))
    return {g: totals[g] / float(denom) for g in totals}

def role_overuse_penalty(
    user_id: int,
    target_group: str,
    rolling_counts: Dict[int, Dict[str, int]],
    rolling_avg: Dict[str, float],
) -> float:
    if not rolling_avg:
        return 1.0
    avg = float(rolling_avg.get(target_group, 0.0))
    count = float(rolling_counts.get(user_id, {}).get(target_group, 0))
    ratio = (count + 1.0) / (avg + 1.0)
    if ratio >= 1.0:
        return max(0.18, 1.0 / (ratio ** 2.2))
    return min(1.35, (1.0 / ratio) ** 0.25)


def _norm_rot_group(g: Optional[str]) -> Optional[str]:
    """
    DB-də bəzən 'NB1/NB2' və ya 'nb1/nb2' kimi saxlanır.
    Rotasiya isə 'nbm/nbm2' gözləyir.
    Bu funksiya hamısını rotasiya formatına çevirir.
    """
    if not g:
        return None
    gg = str(g).strip().lower()

    # NB1/NB2 -> nbm/nbm2
    if gg in {"nb1", "nbm1", "night1"}:
        return "nbm"
    if gg in {"nb2", "nbm2", "night2"}:
        return "nbm2"

    # G1/G2/G3, P1/P2/P3 eyni qalır (lower)
    return gg


def rotation_step_ok(
    last_group: Optional[str],
    target_group: str,
    missing_groups: Optional[list] = None,  # köhnə çağırışlar pozulmasın deyə saxlayırıq
) -> bool:
    """
    Sən istədiyin rotasiya: g3 -> p2 -> g2 -> nbm2 -> p1 -> g1 -> nbm -> p3
    Qayda: user son qrupdan sonra yalnız cycle-də NÖVBƏTİ qrupa keçə bilər (dist == 1).
    """
    lg = _norm_rot_group(last_group)
    tg = _norm_rot_group(target_group)

    if not lg:
        return True  # ilk dəfədirsə sərbəstdir
    if lg not in ROTATION_INDEX or tg not in ROTATION_INDEX:
        return True  # rotasiyaya aid olmayan qrupları bloklama

    cycle_len = len(ROTATION_CYCLE)
    dist = (ROTATION_INDEX[tg] - ROTATION_INDEX[lg]) % cycle_len

    # STRICT: yalnız növbəti addım icazəli (bu ardıcıl NB2-nin qarşısını alır)
    return dist == 1



def group_gap_multiplier(
    user_id: int,
    target_group: str,
    shift_date: date,
    group_last: Optional[Dict[int, Dict[str, date]]],
) -> float:
    if not group_last:
        return 1.0
    last_group_date = group_last.get(user_id, {}).get(target_group)
    if last_group_date is None or shift_date < last_group_date:
        return 1.15
    gap = min((shift_date - last_group_date).days, MAX_GAP_DAYS)
    return 0.8 + (gap / MAX_GAP_DAYS) * 0.6


def same_group_repeat(user: sqlite3.Row, target_group: str, shift_date: date) -> bool:
    ls = parse_date(user["last_selected"])
    if ls is not None and shift_date < ls:
        return False
    lg = user["last_group"]
    return lg == target_group if lg else False


def same_category_repeat(user: sqlite3.Row, target_group: str, shift_date: date) -> bool:
    ls = parse_date(user["last_selected"])
    if ls is not None and shift_date < ls:
        return False
    lg = user["last_group"]
    if not lg:
        return False
    cat_last = group_category(lg)
    cat_target = group_category(target_group)
    if cat_target not in {"gun", "patrul", "nbm"}:
        return False
    return cat_last == cat_target


def group_category(group: str) -> str:
    return GROUP_CATEGORY.get(group, "other")


def build_user_group_history(
    db: sqlite3.Connection,
) -> Tuple[Dict[int, Dict[str, date]], Dict[int, Dict[str, date]]]:
    rows = db.execute(
        """
        SELECT user_id, group_name, MAX(shift_date) AS last_date
        FROM shifts
        GROUP BY user_id, group_name
        """
    ).fetchall()
    group_last: Dict[int, Dict[str, date]] = {}
    category_last: Dict[int, Dict[str, date]] = {}
    for r in rows:
        dt = parse_date(r["last_date"])
        if dt is None:
            continue
        uid = r["user_id"]
        gname = r["group_name"]
        group_last.setdefault(uid, {})[gname] = dt

        cat = group_category(gname)
        category_last.setdefault(uid, {})
        prev = category_last[uid].get(cat)
        if prev is None or dt > prev:
            category_last[uid][cat] = dt
    return group_last, category_last


def group_cooldown_ok(
    user_id: int,
    target_group: str,
    shift_date: date,
    group_last: Dict[int, Dict[str, date]],
    category_last: Dict[int, Dict[str, date]],
    category_cooldown_days: Optional[Dict[str, int]] = None,
    group_cooldown_days: Optional[Dict[str, int]] = None,
) -> bool:
    cat = group_category(target_group)
    cd_map = category_cooldown_days or CATEGORY_COOLDOWN_DAYS
    cd = cd_map.get(cat)
    if cd:
        last_cat = category_last.get(user_id, {}).get(cat)
        if last_cat is not None:
            if last_cat > shift_date:
                pass
            elif last_cat == shift_date:
                pass
            elif (shift_date - last_cat).days < cd:
                return False

    gd_map = group_cooldown_days or GROUP_COOLDOWN_DAYS
    gcd = gd_map.get(target_group)
    if gcd:
        last_group = group_last.get(user_id, {}).get(target_group)
        if last_group is not None:
            if last_group > shift_date:
                pass
            elif last_group == shift_date:
                pass
            elif (shift_date - last_group).days < gcd:
                return False

    return True


def group_penalty_multiplier(
    user_id: int,
    target_group: str,
    shift_date: date,
    group_last: Dict[int, Dict[str, date]],
    category_last: Dict[int, Dict[str, date]],
    category_cooldown_days: Optional[Dict[str, int]] = None,
    group_cooldown_days: Optional[Dict[str, int]] = None,
) -> float:
    mult = 1.0
    cat = group_category(target_group)
    cd_map = category_cooldown_days or CATEGORY_COOLDOWN_DAYS
    cd = cd_map.get(cat)
    if cd:
        last_cat = category_last.get(user_id, {}).get(cat)
        if last_cat is not None:
            if last_cat <= shift_date and (shift_date - last_cat).days < cd:
                mult *= RELAX_PENALTY_CATEGORY

    gd_map = group_cooldown_days or GROUP_COOLDOWN_DAYS
    gcd = gd_map.get(target_group)
    if gcd:
        last_group = group_last.get(user_id, {}).get(target_group)
        if last_group is not None:
            if last_group <= shift_date and (shift_date - last_group).days < gcd:
                mult *= RELAX_PENALTY_GROUP

    return mult


def min_gap_ok(user: sqlite3.Row, shift_date: date, min_gap_days: int = MIN_GAP_DAYS) -> bool:
    ls = parse_date(user["last_selected"])
    if ls is None:
        return True
    if shift_date < ls:
        return True
    if ls == shift_date:
        return True
    return (shift_date - ls).days >= min_gap_days


def user_boluk(user) -> int:
    try:
        return int(user["boluk"])
    except Exception:
        return 1


def user_cycle_mask(user) -> int:
    try:
        val = user["cycle_mask"]
    except Exception:
        try:
            val = user.get("cycle_mask")
        except Exception:
            val = 0
    try:
        return int(val) if val is not None else 0
    except Exception:
        return 0


def cycle_bit(group: str) -> int:
    return ROTATION_BITS.get(group, 0)


def next_cycle_mask(current: int, group: str) -> int:
    bit = cycle_bit(group)
    if not bit:
        return current
    new_mask = current | bit
    if new_mask == ROTATION_FULL_MASK:
        return 0
    return new_mask


def fill_pool_priority(
    pool: List[sqlite3.Row], fill_counts: Dict[int, int]
) -> List[sqlite3.Row]:
    if not pool or not fill_counts:
        return pool
    min_count = min(fill_counts.get(u["id"], 0) for u in pool)
    return [u for u in pool if fill_counts.get(u["id"], 0) == min_count]


def get_boluk2_fill_counts(db: sqlite3.Connection, ym: str) -> Dict[int, int]:
    rows = db.execute(
        "SELECT user_id, count FROM boluk2_fill_history WHERE month=?",
        (ym,),
    ).fetchall()
    return {r["user_id"]: int(r["count"]) for r in rows}


def update_boluk2_fill_history(
    db: sqlite3.Connection, shift_date: date, fill_ids: List[int]
) -> None:
    ds = shift_date.strftime("%Y-%m-%d")
    ym = shift_date.strftime("%Y-%m")
    db.execute("DELETE FROM boluk2_fill_log WHERE shift_date=?", (ds,))
    for uid in fill_ids:
        db.execute(
            """
            INSERT OR IGNORE INTO boluk2_fill_log (month, shift_date, user_id)
            VALUES (?, ?, ?)
            """,
            (ym, ds, uid),
        )
    db.execute("DELETE FROM boluk2_fill_history WHERE month=?", (ym,))
    rows = db.execute(
        """
        SELECT user_id, COUNT(*) AS c
        FROM boluk2_fill_log
        WHERE month=?
        GROUP BY user_id
        """,
        (ym,),
    ).fetchall()
    for r in rows:
        db.execute(
            """
            INSERT INTO boluk2_fill_history (month, user_id, count)
            VALUES (?, ?, ?)
            """,
            (ym, r["user_id"], r["c"]),
        )


def effective_min_gap(user, base_min_gap: int, settings: Optional[Dict[str, object]]) -> int:
    if not settings:
        return base_min_gap
    status_gap = settings.get("status_gap_days") or {}
    try:
        status = normalize_status_code(user["status"])
    except Exception:
        return base_min_gap
    extra = status_gap.get(status)
    try:
        extra_n = int(extra) if extra is not None else 0
    except Exception:
        extra_n = 0
    if extra_n <= 0:
        return base_min_gap
    return max(base_min_gap, extra_n)


def latest_shift_for_user(db: sqlite3.Connection, user_id: int) -> Optional[sqlite3.Row]:
    return db.execute(
        """
        SELECT shift_date, group_name
        FROM shifts
        WHERE user_id=?
        ORDER BY shift_date DESC, id DESC
        LIMIT 1
        """,
        (user_id,),
    ).fetchone()


def update_user_last_selected(db: sqlite3.Connection, user_id: int) -> None:
    latest = latest_shift_for_user(db, user_id)
    if latest:
        db.execute(
            "UPDATE users SET last_selected=?, last_group=? WHERE id=?",
            (latest["shift_date"], latest["group_name"], user_id),
        )
    else:
        db.execute("UPDATE users SET last_selected=NULL, last_group=NULL WHERE id=?", (user_id,))


def replacement_weight_for_bas(
    user: sqlite3.Row, shift_date: date, monthly_counts: Dict[int, int]
) -> float:
    ls = parse_date(user["last_selected"])
    gap = days_between(ls, shift_date)
    w = float(max(0, gap)) ** 2 + 1.0
    mcount = monthly_counts.get(user["id"], 0)
    w *= max(0.25, 1.0 - (mcount * 0.08))
    if gap < 5:
        w *= 0.5
    return max(0.05, w)


def recent_same_group_penalty(user: sqlite3.Row, target_group: str, shift_date: date) -> float:
    """
    "Penalti" = eyni user eyni qrupa tez-tez düşməsin deyə ehtimalı azaldır.
    Əgər son seçilmə 7 gün içindədir və last_group eynidirsə -> weight azaldırıq.
    """
    lg = user["last_group"]
    ls = parse_date(user["last_selected"])
    if ls is not None and shift_date < ls:
        return 1.00
    if lg == target_group and ls is not None and (shift_date - ls).days <= 7:
        return 0.60  # 40% az ehtimal
    return 1.00


def yx_is_eligible(user: sqlite3.Row, shift_date: date) -> bool:
    """
    yx qaydası:
    - həftədə max 1 dəfə -> last_selected ilə min 7 gün
    - yx_next_eligible varsa, shift_date ondan əvvəl olarsa seçilməsin
    """
    if normalize_status_code(user["status"]) != "yx":
        return True

    ls = parse_date(user["last_selected"])
    if ls is not None and shift_date < ls:
        return True
    if ls is not None and (shift_date - ls).days < 7:
        return False

    ne = parse_date(user["yx_next_eligible"])
    if ne is not None and shift_date < ne:
        return False

    return True


def status_allows_any_day(status: str, status_days: Dict[str, List[int]]) -> bool:
    if not status_days:
        return True
    status = normalize_status_code(status)
    if status not in status_days:
        return True
    days = status_days.get(status)
    if days is None:
        return True
    return len(days) > 0


def monthly_quota(
    ym: str,
    users: List[sqlite3.Row],
    status_days: Dict[str, List[int]],
    total_days: Optional[int] = None,
) -> Optional[Dict[str, object]]:
    eligible = [
        u
        for u in users
        if u["role"] == "normal"
        and normalize_status_code(u["status"]) not in ALWAYS_EXCLUDED_STATUSES
        and status_allows_any_day(u["status"], status_days)
    ]
    count = len(eligible)
    if count == 0:
        return None
    if total_days is None:
        start_s, end_s = month_range(ym)
        start = datetime.strptime(start_s, "%Y-%m-%d").date()
        end = datetime.strptime(end_s, "%Y-%m-%d").date()
        total_days = (end - start).days
    total_days = max(1, int(total_days))
    total_slots = total_days * NORMAL_SLOTS
    min_total = total_slots // count
    max_total = min_total + (1 if total_slots % count else 0)
    nb_slots = total_days * (GROUP_SLOT_COUNTS["nbm"] + GROUP_SLOT_COUNTS["nbm2"])
    min_nb = nb_slots // count
    max_nb = min_nb + (1 if nb_slots % count else 0)
    return {
        "eligible_ids": {u["id"] for u in eligible},
        "min_total": min_total,
        "max_total": max_total,
        "min_nb": min_nb,
        "max_nb": max_nb,
    }


def monthly_nb_count(
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]], user_id: int
) -> int:
    if not monthly_group_counts:
        return 0
    counts = monthly_group_counts.get(user_id, {})
    return int(counts.get("nbm", 0)) + int(counts.get("nbm2", 0))


def nb_is_eligible(
    user: sqlite3.Row,
    shift_date: date,
    nb_cooldown_choices: List[int],
    category_last: Optional[Dict[int, Dict[str, date]]],
) -> bool:
    ne = parse_date(user["nb_next_eligible"])
    if ne is not None and shift_date < ne:
        return False
    if not nb_cooldown_choices or not category_last:
        return True
    last_nb = category_last.get(user["id"], {}).get("nbm")
    if last_nb is None or shift_date < last_nb:
        return True
    min_cd = min(nb_cooldown_choices)
    return (shift_date - last_nb).days >= min_cd


def fairness_weight(
    user: sqlite3.Row,
    target_group: str,
    shift_date: date,
    monthly_counts: Dict[int, int],
    prefer_yx: bool,
) -> float:
    """
    Random seçimdir, amma "ədalət" üçün weight (çəki) hesablanır:
    - çoxdan çıxmayanlar daha çox ehtimal
    - dünən çıxanlar daha az ehtimal
    - ayda çox çıxanlar bir az az ehtimal
    - eyni qrup təkrarı az ehtimal
    - yx eligible olanlar (bu gün seçilə bilirsə) bir az üstün (prefer_yx=True olanda)
    """
    ls = parse_date(user["last_selected"])
    gap = days_between(ls, shift_date)

    # base: gap^2 (2 gün çıxmayan çox üstün olur)
    w = float(max(0, gap)) ** 2 + 1.0

    # dünən / bu gün (gap<=1) ehtimal azalsın
    if gap <= 1:
        w *= 0.25

    # aylıq balans
    mcount = monthly_counts.get(user["id"], 0)
    w *= max(0.25, 1.0 - (mcount * 0.08))

    # uzunmüddətli balans (rotation_score)
    w *= max(0.30, 1.0 - (float(user["rotation_score"]) * 0.03))

    # eyni qrup penalti
    w *= recent_same_group_penalty(user, target_group, shift_date)

    # yx eligible-ləri (həftədə 1 dəfə qaydasına görə) bir az üstün tut
    if prefer_yx and normalize_status_code(user["status"]) == "yx" and yx_is_eligible(user, shift_date):
        w *= 1.8

    return max(0.05, w)


def weighted_choice(rows: List[sqlite3.Row], weights: List[float]) -> sqlite3.Row:
    total = sum(weights)
    if total <= 0:
        return _rng.choice(rows)
    r = _rng.random() * total
    upto = 0.0
    for row, w in zip(rows, weights):
        upto += w
        if upto >= r:
            return row
    return rows[-1]


def pick_bas_user(
    candidates: List[dict],
    shift_date: date,
    status_days: Dict[str, List[int]],
    min_gap_days: int,
    settings: Dict[str, object],
    weekend_leave_ids: set,
    monthly_counts: Dict[int, int],
) -> Optional[dict]:
    if not candidates:
        return None
    eligible_bas: List[dict] = []
    eligible_w: List[float] = []
    relaxed_bas: List[dict] = []
    relaxed_w: List[float] = []
    yx_relaxed_bas: List[dict] = []
    yx_relaxed_w: List[float] = []

    for u in candidates:
        if normalize_status_code(u["status"]) in ALWAYS_EXCLUDED_STATUSES:
            continue
        if not status_allowed_today(u["status"], shift_date, status_days):
            continue
        if not weekend_leave_ok(u["id"], shift_date, weekend_leave_ids):
            continue

        yx_ok = yx_is_eligible(u, shift_date)
        ls = parse_date(u["last_selected"])
        gap = days_between(ls, shift_date)
        w = float(max(0, gap)) ** 2 + 1.0
        mcount = monthly_counts.get(u["id"], 0)
        w *= max(0.25, 1.0 - (mcount * 0.08))

        user_min_gap = effective_min_gap(u, min_gap_days, settings)
        if not min_gap_ok(u, shift_date, user_min_gap):
            continue
        if not yx_ok:
            yx_relaxed_bas.append(u)
            yx_relaxed_w.append(max(0.05, w * RELAX_PENALTY_YX))
            continue
        if gap >= 5:
            eligible_bas.append(u)
            eligible_w.append(max(0.05, w))
        else:
            relaxed_bas.append(u)
            relaxed_w.append(max(0.05, w * 0.5))

    if eligible_bas:
        return weighted_choice(eligible_bas, eligible_w)
    if relaxed_bas:
        return weighted_choice(relaxed_bas, relaxed_w)
    if yx_relaxed_bas:
        return weighted_choice(yx_relaxed_bas, yx_relaxed_w)
    return None


def day_eligible_common(
    u: dict,
    shift_date: date,
    status_days: Dict[str, List[int]],
    min_gap_days: int,
    settings: Dict[str, object],
    weekend_leave_ids: set,
) -> bool:
    if normalize_status_code(u["status"]) in ALWAYS_EXCLUDED_STATUSES:
        return False
    if not status_allowed_today(u["status"], shift_date, status_days):
        return False
    if not weekend_leave_ok(u["id"], shift_date, weekend_leave_ids):
        return False
    if not yx_is_eligible(u, shift_date):
        return False
    user_min_gap = effective_min_gap(u, min_gap_days, settings)
    return min_gap_ok(u, shift_date, user_min_gap)


def pick_weighted_random(
    candidates: List[sqlite3.Row],
    target_group: str,
    shift_date: date,
    picks: Dict[str, List[int]],
    monthly_counts: Dict[int, int],
    prefer_yx: bool,
    group_last: Optional[Dict[int, Dict[str, date]]] = None,
    category_last: Optional[Dict[int, Dict[str, date]]] = None,
    monthly_group_counts: Optional[Dict[int, Dict[str, int]]] = None,
    settings: Optional[Dict[str, object]] = None,
    allow_cooldown_break: bool = False,
    allow_rotation_break: bool = False,
) -> sqlite3.Row:
    settings = settings or {}
    status_days = settings.get("status_days") or {}
    min_gap_days = int(settings.get("min_gap_days", MIN_GAP_DAYS))
    category_cooldown_days = settings.get("category_cooldown_days", CATEGORY_COOLDOWN_DAYS)
    group_cooldown_days = settings.get("group_cooldown_days", GROUP_COOLDOWN_DAYS)
    monthly_group_limit = int(settings.get("monthly_group_limit", MONTHLY_GROUP_LIMIT))
    weekend_leave_ids = settings.get("weekend_leave_ids") or set()
    nb_cooldown_choices = settings.get("nb_cooldown_choices", NB_COOLDOWN_CHOICES)
    monthly_quota = settings.get("monthly_quota") or {}
    user_status_by_id = settings.get("user_status_by_id") or {}
    max_yx_per_day = int(settings.get("max_yx_per_day", 0) or 0)
    prev_month_counts = settings.get("prev_month_counts") or {}
    rolling_group_counts = settings.get("rolling_group_counts") or {}
    rolling_group_avg = settings.get("rolling_group_avg") or {}
    quota_ids = monthly_quota.get("eligible_ids") or set()
    quota_min_total = monthly_quota.get("min_total")
    quota_max_total = monthly_quota.get("max_total")
    quota_min_nb = monthly_quota.get("min_nb")
    quota_max_nb = monthly_quota.get("max_nb")

    def collect(
        strict: bool,
        relax_yx: bool,
        relax_group_limit: bool,
        block_repeats: bool,
        allow_cooldown_break: bool,
        allow_rotation_break: bool,
        allow_total_overflow: bool,
        allow_nb_overflow: bool,
        allow_group_repeat: bool,
    ) -> Tuple[List[sqlite3.Row], List[float]]:
        eligible: List[sqlite3.Row] = []
        weights: List[float] = []
        need_total_ids: set = set()
        need_nb_ids: set = set()
        selected_yx_count = 0
        if max_yx_per_day and user_status_by_id:
            for ids in picks.values():
                for uid in ids:
                    if user_status_by_id.get(uid) == "yx":
                        selected_yx_count += 1
        priority_weekend_ids: set[int] = set()
        if weekend_leave_ids and weekday_number(shift_date) == 1:
            for u in candidates:
                if u["id"] not in weekend_leave_ids:
                    continue
                if normalize_status_code(u["status"]) in ALWAYS_EXCLUDED_STATUSES:
                    continue
                if not status_allowed_today(u["status"], shift_date, status_days):
                    continue
                if not yx_is_eligible(u, shift_date):
                    continue
                user_min_gap = effective_min_gap(u, min_gap_days, settings)
                if not min_gap_ok(u, shift_date, user_min_gap):
                    continue
                priority_weekend_ids.add(u["id"])
        selected_priority = 0
        if priority_weekend_ids:
            for ids in picks.values():
                for uid in ids:
                    if uid in priority_weekend_ids:
                        selected_priority += 1
        prev_avg = 0.0
        if prev_month_counts and candidates:
            prev_avg = sum(prev_month_counts.get(u["id"], 0) for u in candidates) / float(
                len(candidates)
            )
        if quota_ids and quota_min_total is not None:
            need_total_ids = {
                u["id"]
                for u in candidates
                if u["id"] in quota_ids and monthly_counts.get(u["id"], 0) < quota_min_total
            }
        if quota_ids and quota_min_nb is not None and target_group in {"nbm", "nbm2"}:
            need_nb_ids = {
                u["id"]
                for u in candidates
                if u["id"] in quota_ids
                and monthly_nb_count(monthly_group_counts, u["id"]) < quota_min_nb
            }
        for u in candidates:
            if normalize_status_code(u["status"]) in ALWAYS_EXCLUDED_STATUSES:
                continue
            if (
                max_yx_per_day
                and selected_yx_count >= max_yx_per_day
                and normalize_status_code(u["status"]) == "yx"
            ):
                continue
            if priority_weekend_ids and selected_priority < len(priority_weekend_ids):
                if u["id"] not in priority_weekend_ids:
                    continue
            if not status_allowed_today(u["status"], shift_date, status_days):
                continue
            if not weekend_leave_ok(u["id"], shift_date, weekend_leave_ids):
                continue
            if already_assigned(u["id"], picks):
                continue
            if need_total_ids and not allow_total_overflow and u["id"] not in need_total_ids:
                continue
            if need_nb_ids and not allow_nb_overflow and u["id"] not in need_nb_ids:
                continue
            missing = (
                user_missing_groups(monthly_group_counts, u["id"])
                if target_group != "bas"
                else []
            )
            if missing and target_group not in missing and not allow_group_repeat:
                continue
            if not allow_group_repeat and cycle_bit(target_group):
                if user_cycle_mask(u) & cycle_bit(target_group):
                    continue

            # yx eligibility rule
            yx_ok = yx_is_eligible(u, shift_date)
            if not relax_yx and not yx_ok:
                continue

            user_min_gap = effective_min_gap(u, min_gap_days, settings)
            if not min_gap_ok(u, shift_date, user_min_gap):
                continue
            if quota_ids and u["id"] in quota_ids:
                total_count = monthly_counts.get(u["id"], 0)
                if quota_max_total is not None and total_count >= quota_max_total and not allow_total_overflow:
                    continue
                if target_group in {"nbm", "nbm2"}:
                    nb_count = monthly_nb_count(monthly_group_counts, u["id"])
                    if quota_max_nb is not None and nb_count >= quota_max_nb and not allow_nb_overflow:
                        continue
            rotation_ok = rotation_step_ok(u["last_group"], target_group, missing)
            if group_last is not None and category_last is not None and not allow_cooldown_break:
                if not rotation_ok and not group_cooldown_ok(
                    u["id"],
                    target_group,
                    shift_date,
                    group_last,
                    category_last,
                    category_cooldown_days,
                    group_cooldown_days,
                ):
                    continue

            if block_repeats:
                if same_group_repeat(u, target_group, shift_date):
                    continue
                if same_category_repeat(u, target_group, shift_date):
                    continue
            if not allow_rotation_break and not rotation_ok:
                continue

            w = fairness_weight(u, target_group, shift_date, monthly_counts, prefer_yx)
            if target_group != "bas":
                if missing:
                    if target_group in missing:
                        w *= 1.35
                    else:
                        w *= RELAX_PENALTY_ROTATION
                if not block_repeats:
                    if same_group_repeat(u, target_group, shift_date):
                        w *= RELAX_PENALTY_SAME_GROUP
                    if same_category_repeat(u, target_group, shift_date):
                        w *= RELAX_PENALTY_SAME_CATEGORY

                gcount = monthly_group_count(monthly_group_counts, u["id"], target_group)
                if gcount >= monthly_group_limit and not relax_group_limit:
                    continue
                w *= monthly_group_factor(gcount)
                if relax_group_limit and gcount >= monthly_group_limit:
                    w *= RELAX_PENALTY_MONTHLY_GROUP
                w *= difficulty_balance_factor(target_group, monthly_group_counts, u["id"])
                w *= transition_multiplier(u, target_group)
                w *= rotation_cycle_multiplier(u, target_group)
                w *= role_overuse_penalty(u["id"], target_group, rolling_group_counts, rolling_group_avg)
                if group_last is not None:
                    w *= group_gap_multiplier(u["id"], target_group, shift_date, group_last)
                if target_group in {"nbm", "nbm2"}:
                    if not nb_is_eligible(u, shift_date, nb_cooldown_choices, category_last):
                        if not allow_cooldown_break:
                            continue
                        w *= RELAX_PENALTY_NBM_DELAY
                if target_group == "nbm":
                    cat_counts = monthly_category_counts(monthly_group_counts, u["id"])
                    if cat_counts["gun"] == 0 or cat_counts["patrul"] == 0:
                        if strict:
                            continue
                        w *= RELAX_PENALTY_NBM_DELAY
                    nbm2_count = monthly_group_count(monthly_group_counts, u["id"], "nbm2")
                    if nbm2_count == 0:
                        if strict:
                            continue
                        w *= RELAX_PENALTY_NBM_DELAY
                    if gcount > 0:
                        w *= 0.6
                if quota_ids and u["id"] in quota_ids:
                    total_count = monthly_counts.get(u["id"], 0)
                    if quota_min_total is not None and total_count < quota_min_total:
                        w *= 1.4
                    if target_group in {"nbm", "nbm2"}:
                        nb_count = monthly_nb_count(monthly_group_counts, u["id"])
                        if quota_min_nb is not None and nb_count < quota_min_nb:
                            w *= 1.4
            if prev_avg > 0:
                prev_count = float(prev_month_counts.get(u["id"], 0))
                if prev_count > prev_avg:
                    w *= max(0.6, 1.0 - (prev_count - prev_avg) * 0.08)
                elif prev_count < prev_avg:
                    w *= min(1.4, 1.0 + (prev_avg - prev_count) * 0.08)
            if weekend_leave_ids and u["id"] in weekend_leave_ids:
                if weekday_number(shift_date) == 1:
                    ls = parse_date(u["last_selected"])
                    if ls is None or (shift_date - ls).days >= 2:
                        w *= 1.25
            if relax_yx and normalize_status_code(u["status"]) == "yx" and not yx_ok:
                w *= RELAX_PENALTY_YX
            if group_last is not None and category_last is not None:
                w *= group_penalty_multiplier(
                    u["id"],
                    target_group,
                    shift_date,
                    group_last,
                    category_last,
                    category_cooldown_days,
                    group_cooldown_days,
                )
            eligible.append(u)
            weights.append(w)
        return eligible, weights

    eligible, weights = collect(
        strict=True,
        relax_yx=False,
        relax_group_limit=False,
        block_repeats=True,
        allow_cooldown_break=False,
        allow_rotation_break=False,
        allow_total_overflow=False,
        allow_nb_overflow=False,
        allow_group_repeat=False,
    )
    if not eligible and group_last is not None and category_last is not None:
        eligible, weights = collect(
            strict=False,
            relax_yx=False,
            relax_group_limit=False,
            block_repeats=True,
            allow_cooldown_break=False,
            allow_rotation_break=False,
            allow_total_overflow=False,
            allow_nb_overflow=False,
            allow_group_repeat=False,
        )
    if not eligible:
        eligible, weights = collect(
            strict=False,
            relax_yx=True,
            relax_group_limit=False,
            block_repeats=True,
            allow_cooldown_break=False,
            allow_rotation_break=False,
            allow_total_overflow=False,
            allow_nb_overflow=False,
            allow_group_repeat=False,
        )
    if not eligible:
        eligible, weights = collect(
            strict=False,
            relax_yx=True,
            relax_group_limit=True,
            block_repeats=True,
            allow_cooldown_break=False,
            allow_rotation_break=False,
            allow_total_overflow=False,
            allow_nb_overflow=False,
            allow_group_repeat=False,
        )
    if not eligible:
        eligible, weights = collect(
            strict=False,
            relax_yx=True,
            relax_group_limit=True,
            block_repeats=True,
            allow_cooldown_break=False,
            allow_rotation_break=False,
            allow_total_overflow=True,
            allow_nb_overflow=True,
            allow_group_repeat=False,
        )
    if not eligible:
        eligible, weights = collect(
            strict=False,
            relax_yx=True,
            relax_group_limit=True,
            block_repeats=True,
            allow_cooldown_break=True,
            allow_rotation_break=False,
            allow_total_overflow=True,
            allow_nb_overflow=True,
            allow_group_repeat=False,
        )
    if not eligible and allow_rotation_break:
        eligible, weights = collect(
            strict=False,
            relax_yx=True,
            relax_group_limit=True,
            block_repeats=True,
            allow_cooldown_break=True,
            allow_rotation_break=True,
            allow_total_overflow=True,
            allow_nb_overflow=True,
            allow_group_repeat=False,
        )
    if not eligible and allow_rotation_break:
        eligible, weights = collect(
            strict=False,
            relax_yx=True,
            relax_group_limit=True,
            block_repeats=False,
            allow_cooldown_break=True,
            allow_rotation_break=True,
            allow_total_overflow=True,
            allow_nb_overflow=True,
            allow_group_repeat=True,
        )

    if not eligible:
        raise RuntimeError("Uyğun user qalmadı. (status/izinli/yx qaydalar?)")

    return weighted_choice(eligible, weights)

def generate_shift(shift_date: date) -> Dict[str, List[int]]:
    db = get_db()
    settings = load_runtime_settings(db)
    status_days = settings["status_days"]
    min_gap_days = settings["min_gap_days"]
    boluk2_enabled = settings.get("boluk2_enabled", False)
    boluk2_cycle_days = settings.get("boluk2_cycle_days", BOLUK2_CYCLE_DAYS)
    ttm_cycle_days = settings.get("ttm_cycle_days", TTM_CYCLE_DAYS)
    last_boluk2_cycle = get_boluk2_last_cycle_before(db, shift_date)
    last_ttm_cycle = get_ttm_last_cycle_before(db, shift_date)
    weekend_leave_ids = get_weekend_leave_ids(db)
    settings["weekend_leave_ids"] = weekend_leave_ids
    picks: Dict[str, List[int]] = {g: [] for g, _ in GROUP_SLOTS}
    picks["bas"] = []

    ym = shift_date.strftime("%Y-%m")
    prev_ym = previous_month(ym)
    monthly_counts = get_monthly_counts(ym)
    monthly_group_counts = get_monthly_group_counts(ym)
    fill_counts = get_boluk2_fill_counts(db, ym)
    group_last, category_last = build_user_group_history(db)
    settings["prev_month_counts"] = get_monthly_counts(prev_ym)
    settings["max_yx_per_day"] = MAX_YX_PER_DAY

    cycle_day = boluk2_enabled and is_boluk2_cycle_day(
        shift_date, last_boluk2_cycle, boluk2_cycle_days
    )
    ttm_cycle_day = is_boluk2_cycle_day(
        shift_date, last_ttm_cycle, ttm_cycle_days
    )

    normal_candidates_all = db.execute("SELECT * FROM users WHERE role='normal'").fetchall()
    normal_by_id = {u["id"]: u for u in normal_candidates_all}
    settings["user_status_by_id"] = {
        u["id"]: normalize_status_code(u["status"]) for u in normal_candidates_all
    }
    eligible_ids = {
        u["id"]
        for u in normal_candidates_all
        if normalize_status_code(u["status"]) not in ALWAYS_EXCLUDED_STATUSES
    }
    rolling_group_counts = get_rolling_group_counts(db, shift_date)
    settings["rolling_group_counts"] = rolling_group_counts
    settings["rolling_group_avg"] = compute_rolling_group_avg(rolling_group_counts, eligible_ids)

    boluk2_active_all = [
        u
        for u in normal_candidates_all
        if u["boluk"] == 2 and normalize_status_code(u["status"]) == "aktiv"
    ]
    effective_cycle_day = cycle_day and bool(boluk2_active_all)

    # ----- BAS (1) -----
    bas_candidates = db.execute("SELECT * FROM users WHERE role='bas'").fetchall()
    if boluk2_enabled:
        preferred_boluk = 2 if effective_cycle_day else 1
        filtered = [u for u in bas_candidates if u["boluk"] == preferred_boluk]
        bas_candidates = filtered or bas_candidates
    else:
        bas_candidates = [u for u in bas_candidates if u["boluk"] == 1]

    chosen_bas = pick_bas_user(
        bas_candidates,
        shift_date,
        status_days,
        min_gap_days,
        settings,
        weekend_leave_ids,
        monthly_counts,
    )
    picks["bas"] = [chosen_bas["id"]] if chosen_bas else []

    # ----- NORMAL (10) -----
    normal_candidates = normal_candidates_all
    forced_ids: set[int] = set()
    fill_needed = 0
    fill_pool: List[sqlite3.Row] = []
    fill_exclude_ids: set[int] = set()
    force_only = False
    if not boluk2_enabled or not effective_cycle_day:
        normal_candidates = [u for u in normal_candidates_all if u["boluk"] == 1]

    ttm_candidates = [
        u for u in normal_candidates if normalize_status_code(u["status"]) == "ttm"
    ]
    ttm_forced_ids: set[int] = set()
    ttm_block_ids: set[int] = set()
    if ttm_candidates:
        if ttm_cycle_day:
            if all(
                day_eligible_common(
                    u, shift_date, status_days, min_gap_days, settings, weekend_leave_ids
                )
                for u in ttm_candidates
            ):
                ttm_forced_ids = {u["id"] for u in ttm_candidates}
            else:
                ttm_block_ids = {u["id"] for u in ttm_candidates}
                normal_candidates = [
                    u for u in normal_candidates if u["id"] not in ttm_block_ids
                ]
        else:
            ttm_block_ids = {u["id"] for u in ttm_candidates}
            normal_candidates = [u for u in normal_candidates if u["id"] not in ttm_block_ids]

    quota = monthly_quota(ym, normal_candidates, status_days)
    if quota:
        settings["monthly_quota"] = quota

    if effective_cycle_day:
        boluk2_active = boluk2_active_all
        forced_ids = {u["id"] for u in boluk2_active}
    forced_ids |= ttm_forced_ids
    if len(forced_ids) > NORMAL_SLOTS and ttm_forced_ids:
        forced_ids -= ttm_forced_ids
        ttm_block_ids |= ttm_forced_ids
        normal_candidates = [u for u in normal_candidates if u["id"] not in ttm_forced_ids]
        ttm_forced_ids = set()

    if len(forced_ids) < NORMAL_SLOTS:
        fill_needed = NORMAL_SLOTS - len(forced_ids)
        if effective_cycle_day:
            fill_pool = [
                u
                for u in normal_candidates_all
                if u["boluk"] == 1 and u["id"] not in ttm_block_ids
            ]
            if last_boluk2_cycle:
                fill_exclude_ids = get_prev_cycle_boluk1_ids(db, last_boluk2_cycle)
        else:
            fill_pool = [u for u in normal_candidates if u["id"] not in ttm_block_ids]
    force_only = bool(forced_ids) and fill_needed == 0

    # yx eligible-ləri (həftədə 1 dəfə qaydasına görə) bir az üstün tut
    for group in GROUP_PICK_ORDER:
        chosen = None
        if forced_ids:
            pool = [u for u in normal_candidates if u["id"] in forced_ids]
            if pool:
                try:
                    chosen = pick_weighted_random(
                        pool,
                        group,
                        shift_date,
                        picks,
                        monthly_counts,
                        prefer_yx=True,
                        group_last=group_last,
                        category_last=category_last,
                        monthly_group_counts=monthly_group_counts,
                        settings=settings,
                        allow_cooldown_break=True,
                        allow_rotation_break=False,
                    )
                except RuntimeError:
                    chosen = None
                if chosen is None:
                    try:
                        chosen = pick_weighted_random(
                            pool,
                            group,
                            shift_date,
                            picks,
                            monthly_counts,
                            prefer_yx=True,
                            group_last=group_last,
                            category_last=category_last,
                            monthly_group_counts=monthly_group_counts,
                            settings=settings,
                            allow_cooldown_break=True,
                            allow_rotation_break=True,
                        )
                    except RuntimeError:
                        chosen = None
            if chosen is not None:
                forced_ids.discard(chosen["id"])
        if chosen is None and fill_needed > 0:
            pool = [u for u in fill_pool if u["id"] not in fill_exclude_ids]
            if not pool:
                pool = fill_pool
            if pool:
                if effective_cycle_day:
                    pool = fill_pool_priority(pool, fill_counts)
                try:
                    chosen = pick_weighted_random(
                        pool,
                        group,
                        shift_date,
                        picks,
                        monthly_counts,
                        prefer_yx=True,
                        group_last=group_last,
                        category_last=category_last,
                        monthly_group_counts=monthly_group_counts,
                        settings=settings,
                    )
                except RuntimeError:
                    try:
                        chosen = pick_weighted_random(
                            pool,
                            group,
                            shift_date,
                            picks,
                            monthly_counts,
                            prefer_yx=True,
                            group_last=group_last,
                            category_last=category_last,
                            monthly_group_counts=monthly_group_counts,
                            settings=settings,
                            allow_cooldown_break=True,
                            allow_rotation_break=False,
                        )
                    except RuntimeError:
                        chosen = None
                if chosen is None:
                    try:
                        chosen = pick_weighted_random(
                            pool,
                            group,
                            shift_date,
                            picks,
                            monthly_counts,
                            prefer_yx=True,
                            group_last=group_last,
                            category_last=category_last,
                            monthly_group_counts=monthly_group_counts,
                            settings=settings,
                            allow_cooldown_break=True,
                            allow_rotation_break=True,
                        )
                    except RuntimeError:
                        chosen = None
            if chosen is not None:
                fill_needed -= 1
        if chosen is None and not force_only:
            try:
                chosen = pick_weighted_random(
                    normal_candidates,
                    group,
                    shift_date,
                    picks,
                    monthly_counts,
                    prefer_yx=True,
                    group_last=group_last,
                    category_last=category_last,
                    monthly_group_counts=monthly_group_counts,
                    settings=settings,
                )
            except RuntimeError:
                try:
                    chosen = pick_weighted_random(
                        normal_candidates,
                        group,
                        shift_date,
                        picks,
                        monthly_counts,
                        prefer_yx=True,
                        group_last=group_last,
                        category_last=category_last,
                        monthly_group_counts=monthly_group_counts,
                        settings=settings,
                        allow_cooldown_break=True,
                        allow_rotation_break=False,
                    )
                except RuntimeError:
                    chosen = None
            if chosen is None:
                try:
                    chosen = pick_weighted_random(
                        normal_candidates,
                        group,
                        shift_date,
                        picks,
                        monthly_counts,
                        prefer_yx=True,
                        group_last=group_last,
                        category_last=category_last,
                        monthly_group_counts=monthly_group_counts,
                        settings=settings,
                        allow_cooldown_break=True,
                        allow_rotation_break=True,
                    )
                except RuntimeError:
                    continue
        if chosen is not None and forced_ids and chosen["id"] in forced_ids:
            forced_ids.discard(chosen["id"])
        picks[group].append(chosen["id"])

    if effective_cycle_day or ttm_forced_ids:
        selected_ids = {uid for ids in picks.values() for uid in ids}
        forced_expected: set[int] = set(ttm_forced_ids)
        if effective_cycle_day:
            forced_expected |= {u["id"] for u in boluk2_active}
        missing_forced = forced_expected - selected_ids
        if missing_forced:
            users_by_id = {u["id"]: u for u in normal_candidates_all}
            swap_in_forced_users(picks, missing_forced, users_by_id, monthly_group_counts)
            selected_ids = {uid for ids in picks.values() for uid in ids}
    return picks


def save_shift(shift_date: date, picks: Dict[str, List[int]]):
    db = get_db()
    settings = load_runtime_settings(db)
    yx_cooldown_choices = settings["yx_cooldown_choices"]
    nb_cooldown_choices = settings.get("nb_cooldown_choices", NB_COOLDOWN_CHOICES)
    boluk2_enabled = settings.get("boluk2_enabled", False)
    boluk2_cycle_days = settings.get("boluk2_cycle_days", BOLUK2_CYCLE_DAYS)
    ttm_cycle_days = settings.get("ttm_cycle_days", TTM_CYCLE_DAYS)
    ds = shift_date.strftime("%Y-%m-%d")

    # overwrite
    db.execute("DELETE FROM shifts WHERE shift_date=?", (ds,))

    for group, ids in picks.items():
        for uid in ids:
            db.execute(
                "INSERT OR IGNORE INTO shifts (shift_date, group_name, user_id) VALUES (?, ?, ?)",
                (ds, group, uid),
            )

    selected = [(uid, group) for group, ids in picks.items() for uid in ids]
    for uid, group in selected:
        # yx eligible-ləri (həftədə 1 dəfə qaydasına görə) bir az üstün tut
        user = db.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
        next_eligible = None
        next_nb_eligible = None
        cycle_started = user["cycle_started"] if user else None
        current_mask = user_cycle_mask(user) if user else 0
        new_mask = next_cycle_mask(current_mask, group)
        if user and normalize_status_code(user["status"]) == "yx":
            cooldown = _rng.choice(yx_cooldown_choices)
            next_eligible = (shift_date + timedelta(days=cooldown)).strftime("%Y-%m-%d")
        if group in {"nbm", "nbm2"}:
            cooldown = _rng.choice(nb_cooldown_choices)
            next_nb_eligible = (shift_date + timedelta(days=cooldown)).strftime("%Y-%m-%d")
        if current_mask == 0 and new_mask != 0 and not cycle_started:
            cycle_started = ds
        if current_mask != 0 and new_mask == 0:
            cycle_started = ds

        db.execute(
            """
            UPDATE users
            SET last_selected=?,
                rotation_score=rotation_score + 1,
                last_group=?,
                yx_next_eligible=COALESCE(?, yx_next_eligible),
                nb_next_eligible=COALESCE(?, nb_next_eligible),
                cycle_mask=?,
                cycle_started=?
            WHERE id=?
            """,
            (ds, group, next_eligible, next_nb_eligible, new_mask, cycle_started, uid),
        )

    # Small decay for long-term fairness
    db.execute(
        """
        UPDATE users
        SET rotation_score = CASE
            WHEN rotation_score > 0 THEN rotation_score - 0.2
            ELSE rotation_score
        END
        WHERE role='normal'
        """
    )

    if boluk2_enabled:
        last_cycle = get_boluk2_last_cycle(db)
        if is_boluk2_cycle_day(shift_date, last_cycle, boluk2_cycle_days):
            boluk2_active = db.execute(
                """
                SELECT id FROM users
                WHERE role='normal' AND boluk=2 AND LOWER(TRIM(status))='aktiv'
                """
            ).fetchall()
            boluk2_active_ids = {r["id"] for r in boluk2_active}
            if boluk2_active_ids:
                selected_ids = {uid for group, ids in picks.items() if group != "bas" for uid in ids}
                boluk2_ok = boluk2_active_ids.issubset(selected_ids)
                if boluk2_ok:
                    set_setting(db, "boluk2_last_cycle", shift_date.strftime("%Y-%m-%d"))
                fill_ids: List[int] = []
                if selected_ids:
                    rows = db.execute(
                        """
                        SELECT id, boluk
                        FROM users
                        WHERE role='normal' AND id IN ({})
                        """.format(
                            ",".join("?" * len(selected_ids))
                        ),
                        tuple(selected_ids),
                    ).fetchall()
                    fill_ids = [r["id"] for r in rows if int(r["boluk"]) == 1]
                if not boluk2_ok:
                    fill_ids = []
                update_boluk2_fill_history(db, shift_date, fill_ids)

    last_ttm_cycle = parse_date(get_setting(db, "ttm_last_cycle", None))
    if is_boluk2_cycle_day(shift_date, last_ttm_cycle, ttm_cycle_days):
        ttm_active = db.execute(
            """
            SELECT id FROM users
            WHERE role='normal' AND LOWER(TRIM(status)) IN ('ttm', 'tk')
            """
        ).fetchall()
        ttm_active_ids = {r["id"] for r in ttm_active}
        if ttm_active_ids:
            selected_ids = {
                uid for group, ids in picks.items() if group != "bas" for uid in ids
            }
            if ttm_active_ids.issubset(selected_ids):
                set_setting(db, "ttm_last_cycle", shift_date.strftime("%Y-%m-%d"))

    db.commit()


def load_shift(shift_date: date) -> Dict[str, List[sqlite3.Row]]:
    db = get_db()
    settings = load_runtime_settings(db)
    boluk2_enabled = settings.get("boluk2_enabled", False)
    weekend_leave_ids = get_weekend_leave_ids(db)
    ds = shift_date.strftime("%Y-%m-%d")
    rows = db.execute(
        """
        SELECT s.group_name, u.*
        FROM shifts s
        JOIN users u ON u.id = s.user_id
        WHERE s.shift_date=?
        ORDER BY s.group_name, s.id
        """,
        (ds,),
    ).fetchall()

    out: Dict[str, List[sqlite3.Row]] = {}
    for r in rows:
        out.setdefault(r["group_name"], []).append(r)
    return out


def get_shift_confirmed_at(db: sqlite3.Connection, ds: str) -> Optional[str]:
    row = db.execute(
        "SELECT confirmed_at FROM shift_meta WHERE shift_date=?", (ds,)
    ).fetchone()
    return row["confirmed_at"] if row else None


def set_shift_confirmed(db: sqlite3.Connection, ds: str, confirmed_at: Optional[str]) -> None:
    db.execute(
        """
        INSERT INTO shift_meta (shift_date, confirmed_at)
        VALUES (?, ?)
        ON CONFLICT(shift_date) DO UPDATE SET confirmed_at=excluded.confirmed_at
        """,
        (ds, confirmed_at),
    )


def materialize_picks(picks: Dict[str, List[int]]) -> Dict[str, List[sqlite3.Row]]:
    db = get_db()
    user_ids = [uid for ids in picks.values() for uid in ids]
    if not user_ids:
        return {}
    placeholders = ",".join(["?"] * len(user_ids))
    rows = db.execute(
        f"SELECT * FROM users WHERE id IN ({placeholders})", user_ids
    ).fetchall()
    by_id = {r["id"]: r for r in rows}
    out: Dict[str, List[sqlite3.Row]] = {}
    for group, ids in picks.items():
        out[group] = [by_id[uid] for uid in ids if uid in by_id]
    return out


def get_user_history(user_id: int, ym: str) -> List[sqlite3.Row]:
    db = get_db()
    start, end = month_range(ym)
    return db.execute(
        """
        SELECT s.shift_date, s.group_name
        FROM shifts s
        LEFT JOIN shift_meta sm ON sm.shift_date = s.shift_date
        WHERE s.user_id = ?
          AND s.shift_date >= ? AND s.shift_date < ?
          AND (sm.confirmed_at IS NOT NULL OR sm.shift_date IS NULL)
        ORDER BY s.shift_date ASC, s.group_name ASC
        """,
        (user_id, start, end),
    ).fetchall()


@app.route("/")
def home():
    # default shift date = today (shift starts 18:00)
    return redirect(url_for("naryad", d=date.today().strftime("%Y-%m-%d")))


@app.route("/naryad")
def naryad():
    d = request.args.get("d") or date.today().strftime("%Y-%m-%d")
    preview = False
    shift_date = datetime.strptime(d, "%Y-%m-%d").date()
    shift_next_date = shift_date + timedelta(days=1)
    ds = shift_date.strftime("%Y-%m-%d")
    data = load_shift(shift_date)
    db = get_db()
    settings = load_runtime_settings(db)
    status_days = settings["status_days"]
    min_gap_days = settings["min_gap_days"]
    boluk2_enabled = settings.get("boluk2_enabled", False)
    weekend_leave_ids = get_weekend_leave_ids(db)
    users = db.execute("SELECT * FROM users ORDER BY id ASC").fetchall()
    confirmed_at = get_shift_confirmed_at(db, ds)

    selected_ids = {r["id"] for group in data.values() for r in group}

    status_order = ["yx", "ttm", "tm", "tp", "izinli", "aktiv"]
    status_groups = {s: [] for s in status_order}
    other_status = []
    for u in users:
        status = normalize_status_code(u["status"])
        if status in status_groups:
            status_groups[status].append(u)
        else:
            other_status.append(u)
    if other_status:
        status_groups["other"] = other_status
        status_order.append("other")

    free_normals = [
        u
        for u in users
        if u["role"] == "normal"
        and u["status"] == "aktiv"
        and u["id"] not in selected_ids
    ]
    free_boluk1 = [u for u in free_normals if u["boluk"] == 1]
    free_boluk2 = [u for u in free_normals if u["boluk"] == 2]
    boluk1_users = [u for u in users if u["boluk"] == 1]
    boluk2_users = [u for u in users if u["boluk"] == 2]

    ym = shift_date.strftime("%Y-%m")
    monthly_counts = get_monthly_counts(ym)
    bas_candidates = [u for u in users if u["role"] == "bas"]
    normal_candidates = [u for u in users if u["role"] == "normal"]

    replacement_options: Dict[str, Dict[int, List[sqlite3.Row]]] = {}
    if not preview:
        for group, rows in data.items():
            if not rows:
                continue
            is_bas = group == "bas"
            candidates = bas_candidates if is_bas else normal_candidates
            for current in rows:
                options: List[sqlite3.Row] = []
                for u in candidates:
                    if u["id"] in selected_ids:
                        continue
                    options.append(u)
                options.sort(key=lambda row: str(row["name"]).lower())
                replacement_options.setdefault(group, {})
                replacement_options[group][current["id"]] = options

    change_map: Dict[str, Dict[int, sqlite3.Row]] = {}
    if not preview:
        changes = db.execute(
            """
            SELECT sc.*, uo.name AS old_name, un.name AS new_name
            FROM shift_changes sc
            LEFT JOIN users uo ON uo.id = sc.old_user_id
            LEFT JOIN users un ON un.id = sc.new_user_id
            WHERE sc.shift_date=?
            ORDER BY sc.created_at DESC, sc.id DESC
            """,
            (ds,),
        ).fetchall()
        for c in changes:
            gname = c["group_name"]
            change_map.setdefault(gname, {})
            if c["new_user_id"] not in change_map[gname]:
                change_map[gname][c["new_user_id"]] = c

    swap_candidates: Dict[int, List[Dict[str, object]]] = {}
    if not preview:
        selected_all: List[Dict[str, object]] = []
        for group, rows in data.items():
            for u in rows:
                selected_all.append(
                    {
                        "id": u["id"],
                        "name": u["name"],
                        "group": group,
                        "role": u["role"],
                    }
                )
        for u in selected_all:
            swap_candidates[u["id"]] = [
                cand
                for cand in selected_all
                if cand["id"] != u["id"]
                and cand["role"] == u["role"]
                and cand["group"] != u["group"]
            ]

    return render_template(
        "naryad.html",
        shift_date=shift_date,
        shift_next_date=shift_next_date,
        data=data,
        users=users,
        is_preview=preview,
        shift_confirmed_at=confirmed_at,
        has_selection=bool(selected_ids),
        boluk2_enabled=boluk2_enabled,
        boluk1_users=boluk1_users,
        boluk2_users=boluk2_users,
        status_order=status_order,
        status_groups=status_groups,
        free_normals=free_normals,
        free_boluk1=free_boluk1,
        free_boluk2=free_boluk2,
        replacement_options=replacement_options,
        change_map=change_map,
        swap_candidates=swap_candidates,
        shift_hour=SHIFT_START_HOUR,
    )


@app.route("/toggle_boluk2", methods=["POST"])
def toggle_boluk2():
    d = request.form.get("d") or date.today().strftime("%Y-%m-%d")
    enabled_raw = (request.form.get("enabled") or "").strip().lower()
    enabled = enabled_raw in {"1", "true", "on", "yes"}
    db = get_db()
    set_setting(db, "boluk2_enabled", enabled)
    msg = "2-ci boluk aktivdir." if enabled else "2-ci boluk sonduruldu."
    flash(msg, "ok")
    return redirect(url_for("naryad", d=d))


@app.route("/generate", methods=["POST"])
def generate():
    d = request.form.get("d") or date.today().strftime("%Y-%m-%d")
    shift_date = datetime.strptime(d, "%Y-%m-%d").date()
    picks = generate_shift(shift_date)
    save_shift(shift_date, picks)
    db = get_db()
    settings = load_runtime_settings(db)
    yx_cooldown_choices = settings["yx_cooldown_choices"]
    ds = shift_date.strftime("%Y-%m-%d")
    set_shift_confirmed(db, ds, None)
    db.commit()
    flash("Növbə yaradıldı (18:00 -> növbəti gün 18:00).", "ok")
    return redirect(url_for("naryad", d=shift_date.strftime("%Y-%m-%d")))


@app.route("/confirm_shift", methods=["POST"])
def confirm_shift():
    d = request.form.get("d") or date.today().strftime("%Y-%m-%d")
    try:
        shift_date = datetime.strptime(d, "%Y-%m-%d").date()
    except ValueError:
        flash("Invalid date.", "err")
        return redirect(url_for("naryad", d=d))

    db = get_db()
    ds = shift_date.strftime("%Y-%m-%d")
    count = db.execute(
        "SELECT COUNT(*) AS c FROM shifts WHERE shift_date=?", (ds,)
    ).fetchone()["c"]
    if count <= 0:
        flash("Təsdiqləmək üçün naryad yoxdur.", "err")
        return redirect(url_for("naryad", d=ds))

    confirmed_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    set_shift_confirmed(db, ds, confirmed_at)
    db.commit()
    flash("Naryad təsdiqləndi.", "ok")
    return redirect(url_for("naryad", d=ds))


@app.route("/clear_shift", methods=["POST"])
def clear_shift():
    d = request.form.get("d") or date.today().strftime("%Y-%m-%d")
    try:
        shift_date = datetime.strptime(d, "%Y-%m-%d").date()
    except ValueError:
        flash("Invalid date.", "err")
        return redirect(url_for("naryad", d=d))

    db = get_db()
    ds = shift_date.strftime("%Y-%m-%d")
    affected = db.execute(
        "SELECT DISTINCT user_id FROM shifts WHERE shift_date=?", (ds,)
    ).fetchall()

    db.execute("DELETE FROM shifts WHERE shift_date=?", (ds,))
    db.execute("DELETE FROM shift_changes WHERE shift_date=?", (ds,))
    db.execute("DELETE FROM shift_meta WHERE shift_date=?", (ds,))

    for r in affected:
        uid = r["user_id"]
        db.execute(
            """
            UPDATE users
            SET rotation_score = CASE
                WHEN rotation_score > 0 THEN rotation_score - 1
                ELSE 0
            END
            WHERE id=?
            """,
            (uid,),
        )
        update_user_last_selected(db, uid)
        user_state = db.execute(
            "SELECT status, last_selected FROM users WHERE id=?", (uid,)
        ).fetchone()
        if user_state and user_state["status"] == "yx":
            ls = parse_date(user_state["last_selected"])
            if ls is None or ls < shift_date:
                db.execute(
                    "UPDATE users SET yx_next_eligible=NULL WHERE id=?", (uid,)
                )

    db.commit()
    flash("Seçim təmizləndi.", "ok")
    return redirect(url_for("naryad", d=ds))


@app.route("/clear_all", methods=["POST"])
def clear_all():
    d = request.form.get("d") or date.today().strftime("%Y-%m-%d")
    db = get_db()
    db.execute("DELETE FROM shifts")
    db.execute("DELETE FROM shift_changes")
    db.execute("DELETE FROM shift_meta")
    db.execute(
        """
        UPDATE users
        SET last_selected=NULL,
            last_group=NULL,
            rotation_score=0,
            yx_next_eligible=NULL
        """
    )
    db.commit()
    flash("Bütün naryad tarixçəsi təmizləndi.", "ok")
    return redirect(url_for("naryad", d=d))


@app.route("/users", methods=["GET"])
def users():
    ym = request.args.get("month") or date.today().strftime("%Y-%m")
    db = get_db()
    users = db.execute("SELECT * FROM users ORDER BY id ASC").fetchall()
    boluk1_users = [u for u in users if u["boluk"] == 1]
    boluk2_users = [u for u in users if u["boluk"] == 2]
    monthly_counts = get_monthly_counts(ym)
    return render_template(
        "users.html",
        users=users,
        boluk1_users=boluk1_users,
        boluk2_users=boluk2_users,
        ym=ym,
        monthly_counts=monthly_counts,
    )


@app.route("/buraxlis", methods=["GET", "POST"])
def buraxlis():
    db = get_db()
    users = db.execute("SELECT * FROM users ORDER BY id ASC").fetchall()
    boluk1_users = [u for u in users if u["boluk"] == 1]
    boluk2_users = [u for u in users if u["boluk"] == 2]
    selected_ids = get_weekend_leave_ids(db)
    if request.method == "POST":
        raw_ids = request.form.getlist("user_id")
        ids: List[int] = []
        for raw in raw_ids:
            try:
                ids.append(int(raw))
            except Exception:
                continue
        ids = list(dict.fromkeys(ids))
        db.execute("DELETE FROM weekend_leave")
        for uid in ids:
            db.execute(
                "INSERT OR REPLACE INTO weekend_leave (user_id, enabled) VALUES (?, 1)",
                (uid,),
            )
        db.commit()
        flash("Buraxlis yadda saxlandi.", "ok")
        return redirect(url_for("buraxlis"))
    return render_template(
        "buraxlis.html",
        users=users,
        boluk1_users=boluk1_users,
        boluk2_users=boluk2_users,
        selected_ids=selected_ids,
    )


@app.route("/users/add", methods=["POST"])
def users_add():
    name = (request.form.get("name") or "").strip()
    role = request.form.get("role") or "normal"
    status = normalize_status_code(request.form.get("status") or "aktiv")
    boluk_raw = request.form.get("boluk") or "1"
    ym = request.form.get("month") or date.today().strftime("%Y-%m")
    try:
        boluk = 2 if int(boluk_raw) == 2 else 1
    except Exception:
        boluk = 1

    if role not in VALID_ROLES or status not in VALID_STATUSES:
        flash("Invalid role/status.", "err")
        return redirect(url_for("users", month=ym))

    db = get_db()
    max_id = db.execute("SELECT COALESCE(MAX(id), 0) AS m FROM users").fetchone()["m"]
    new_id = int(max_id) + 1
    if not name:
        name = f"User {new_id}"
    db.execute(
        "INSERT INTO users (id, name, role, status, boluk) VALUES (?, ?, ?, ?, ?)",
        (new_id, name, role, status, boluk),
    )
    db.commit()
    flash("User added.", "ok")
    return redirect(url_for("users", month=ym))


@app.route("/users/delete", methods=["POST"])
def users_delete():
    uid = request.form.get("delete_id", type=int)
    ym = request.form.get("month") or date.today().strftime("%Y-%m")
    if not uid:
        flash("User id missing.", "err")
        return redirect(url_for("users", month=ym))

    db = get_db()
    db.execute("DELETE FROM shifts WHERE user_id=?", (uid,))
    db.execute("DELETE FROM shift_changes WHERE old_user_id=? OR new_user_id=?", (uid, uid))
    db.execute("DELETE FROM weekend_leave WHERE user_id=?", (uid,))
    db.execute("DELETE FROM users WHERE id=?", (uid,))
    db.commit()
    flash("User deleted.", "ok")
    return redirect(url_for("users", month=ym))


@app.route("/users/save_all", methods=["POST"])
def users_save_all():
    db = get_db()
    ids = request.form.getlist("id[]")
    names = request.form.getlist("name[]")
    roles = request.form.getlist("role[]")
    statuses = request.form.getlist("status[]")
    boluks = request.form.getlist("boluk[]")

    if not (len(ids) == len(names) == len(roles) == len(statuses) == len(boluks)):
        return "Form data error", 400

    for i in range(len(ids)):
        uid = int(ids[i])
        name = (names[i] or "").strip() or f"User {uid}"
        role = roles[i]
        status = normalize_status_code(statuses[i])
        try:
            boluk = 2 if int(boluks[i]) == 2 else 1
        except Exception:
            boluk = 1
        db.execute(
            "UPDATE users SET name=?, role=?, status=?, boluk=? WHERE id=?",
            (name, role, status, boluk, uid),
        )

        # yx eligible-ləri (həftədə 1 dəfə qaydasına görə) bir az üstün tut
        if status != "yx":
            db.execute("UPDATE users SET yx_next_eligible=NULL WHERE id=?", (uid,))

    db.commit()
    flash("Bütün user-lər save olundu.", "ok")

    ym = request.form.get("month") or date.today().strftime("%Y-%m")
    return redirect(url_for("users", month=ym))


@app.route("/change_user", methods=["POST"])
def change_user():
    d = request.form.get("d") or date.today().strftime("%Y-%m-%d")
    group = (request.form.get("group") or "").strip()
    old_user_id = request.form.get("old_user_id", type=int)
    new_user_id = request.form.get("new_user_id", type=int)
    swap_user_id = request.form.get("swap_user_id", type=int)
    reason = (request.form.get("reason") or "").strip()

    try:
        shift_date = datetime.strptime(d, "%Y-%m-%d").date()
    except ValueError:
        flash("Invalid date.", "err")
        return redirect(url_for("naryad", d=d))

    if not group or not old_user_id or not new_user_id:
        if swap_user_id:
            pass
        else:
            flash("Missing change data.", "err")
            return redirect(url_for("naryad", d=shift_date.strftime("%Y-%m-%d")))

    db = get_db()
    ds = shift_date.strftime("%Y-%m-%d")

    existing = db.execute(
        """
        SELECT 1
        FROM shifts
        WHERE shift_date=? AND group_name=? AND user_id=?
        """,
        (ds, group, old_user_id),
    ).fetchone()
    if not existing:
        flash("Selected user not found in this shift.", "err")
        return redirect(url_for("naryad", d=ds))

    if swap_user_id:
        if swap_user_id == old_user_id:
            flash("Swap user is the same.", "err")
            return redirect(url_for("naryad", d=ds))

        row_old = db.execute(
            "SELECT group_name FROM shifts WHERE shift_date=? AND user_id=?",
            (ds, old_user_id),
        ).fetchone()
        row_swap = db.execute(
            "SELECT group_name FROM shifts WHERE shift_date=? AND user_id=?",
            (ds, swap_user_id),
        ).fetchone()
        if not row_swap or not row_old:
            flash("Swap user not found in this shift.", "err")
            return redirect(url_for("naryad", d=ds))

        group_old = row_old["group_name"]
        group_swap = row_swap["group_name"]
        if group_old == group_swap:
            flash("Both users are in the same group.", "err")
            return redirect(url_for("naryad", d=ds))

        old_user = db.execute("SELECT * FROM users WHERE id=?", (old_user_id,)).fetchone()
        swap_user = db.execute("SELECT * FROM users WHERE id=?", (swap_user_id,)).fetchone()
        if not old_user or not swap_user:
            flash("Swap users not found.", "err")
            return redirect(url_for("naryad", d=ds))

        if group_old == "bas" or group_swap == "bas":
            if old_user["role"] != "bas" or swap_user["role"] != "bas":
                flash("Bas role cannot be swapped with normal.", "err")
                return redirect(url_for("naryad", d=ds))
        else:
            if old_user["role"] != "normal" or swap_user["role"] != "normal":
                flash("Only normal users can be swapped in normal groups.", "err")
                return redirect(url_for("naryad", d=ds))

        # Swap in shift table
        db.execute(
            "DELETE FROM shifts WHERE shift_date=? AND group_name=? AND user_id=?",
            (ds, group_old, old_user_id),
        )
        db.execute(
            "DELETE FROM shifts WHERE shift_date=? AND group_name=? AND user_id=?",
            (ds, group_swap, swap_user_id),
        )
        db.execute(
            "INSERT INTO shifts (shift_date, group_name, user_id) VALUES (?, ?, ?)",
            (ds, group_old, swap_user_id),
        )
        db.execute(
            "INSERT INTO shifts (shift_date, group_name, user_id) VALUES (?, ?, ?)",
            (ds, group_swap, old_user_id),
        )

        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        reason_a = reason or f"Swap with {swap_user_id}"
        reason_b = reason or f"Swap with {old_user_id}"
        db.execute(
            """
            INSERT INTO shift_changes (shift_date, group_name, old_user_id, new_user_id, reason, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (ds, group_old, old_user_id, swap_user_id, reason_a[:200], ts),
        )
        db.execute(
            """
            INSERT INTO shift_changes (shift_date, group_name, old_user_id, new_user_id, reason, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (ds, group_swap, swap_user_id, old_user_id, reason_b[:200], ts),
        )

        update_user_last_selected(db, old_user_id)
        update_user_last_selected(db, swap_user_id)
        set_shift_confirmed(db, ds, None)
        db.commit()
        flash("Swap applied.", "ok")
        return redirect(url_for("naryad", d=ds))

    already_in_shift = db.execute(
        "SELECT 1 FROM shifts WHERE shift_date=? AND user_id=?",
        (ds, new_user_id),
    ).fetchone()
    if already_in_shift:
        flash("New user is already in this shift.", "err")
        return redirect(url_for("naryad", d=ds))

    new_user = db.execute("SELECT * FROM users WHERE id=?", (new_user_id,)).fetchone()
    if not new_user:
        flash("New user not found.", "err")
        return redirect(url_for("naryad", d=ds))

    old_user_full = db.execute("SELECT * FROM users WHERE id=?", (old_user_id,)).fetchone()
    if not old_user_full:
        flash("Old user not found.", "err")
        return redirect(url_for("naryad", d=ds))

    required_role = "bas" if group == "bas" else "normal"
    if new_user["role"] != required_role:
        flash(f"New user must be role '{required_role}'.", "err")
        return redirect(url_for("naryad", d=ds))

    # Swap in shift table
    db.execute(
        "DELETE FROM shifts WHERE shift_date=? AND group_name=? AND user_id=?",
        (ds, group, old_user_id),
    )
    db.execute(
        "INSERT INTO shifts (shift_date, group_name, user_id) VALUES (?, ?, ?)",
        (ds, group, new_user_id),
    )
    db.execute(
        """
        INSERT INTO shift_changes (shift_date, group_name, old_user_id, new_user_id, reason, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            ds,
            group,
            old_user_id,
            new_user_id,
            reason[:200] if reason else None,
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        ),
    )

    # Adjust stats without reapplying global decay
    db.execute(
        """
        UPDATE users
        SET rotation_score = CASE
            WHEN rotation_score > 0 THEN rotation_score - 1
            ELSE 0
        END
        WHERE id=?
        """,
        (old_user_id,),
    )
    db.execute(
        "UPDATE users SET rotation_score = rotation_score + 1 WHERE id=?",
        (new_user_id,),
    )
    update_user_last_selected(db, old_user_id)
    update_user_last_selected(db, new_user_id)

    old_user = db.execute("SELECT status, last_selected FROM users WHERE id=?", (old_user_id,)).fetchone()
    if old_user and normalize_status_code(old_user["status"]) == "yx":
        ls = parse_date(old_user["last_selected"])
        if ls is None or ls < shift_date:
            db.execute("UPDATE users SET yx_next_eligible=NULL WHERE id=?", (old_user_id,))

    new_user_state = db.execute(
        "SELECT status, last_selected FROM users WHERE id=?", (new_user_id,)
    ).fetchone()
    if new_user_state and normalize_status_code(new_user_state["status"]) == "yx":
        ls = parse_date(new_user_state["last_selected"])
        if ls == shift_date:
            cooldown = _rng.choice(yx_cooldown_choices)
            next_eligible = (shift_date + timedelta(days=cooldown)).strftime("%Y-%m-%d")
            db.execute(
                "UPDATE users SET yx_next_eligible=? WHERE id=?",
                (next_eligible, new_user_id),
            )

    set_shift_confirmed(db, ds, None)
    db.commit()
    flash("Naryad yeniləndi.", "ok")
    return redirect(url_for("naryad", d=ds))


@app.route("/history")
def history():
    ym = request.args.get("month") or date.today().strftime("%Y-%m")
    user_id_raw = request.args.get("user_id")
    show_all = user_id_raw == "all"
    boluk_raw = (request.args.get("boluk") or "").strip()
    boluk_filter = None
    if boluk_raw in {"1", "2"}:
        boluk_filter = int(boluk_raw)
    user_id = None
    if user_id_raw and not show_all:
        try:
            user_id = int(user_id_raw)
        except ValueError:
            user_id = None

    start_s, end_s = month_range(ym)
    start = datetime.strptime(start_s, "%Y-%m-%d").date()
    end = datetime.strptime(end_s, "%Y-%m-%d").date()
    total_days = (end - start).days
    days = [start + timedelta(days=i) for i in range(total_days)]

    db = get_db()
    if boluk_filter:
        users = db.execute(
            "SELECT id, name, boluk FROM users WHERE boluk=? ORDER BY id ASC",
            (boluk_filter,),
        ).fetchall()
    else:
        users = db.execute("SELECT id, name, boluk FROM users ORDER BY id ASC").fetchall()

    history_rows = []
    all_rows = []
    selected_user = None
    matrix_users = []
    matrix: Dict[int, Dict[str, str]] = {}
    if show_all:
        if boluk_filter:
            all_rows = db.execute(
                """
                SELECT s.shift_date, s.group_name, u.id AS user_id, u.name AS user_name
                FROM shifts s
                JOIN users u ON u.id = s.user_id
                LEFT JOIN shift_meta sm ON sm.shift_date = s.shift_date
                WHERE s.shift_date >= ? AND s.shift_date < ?
                  AND u.boluk = ?
                  AND (sm.confirmed_at IS NOT NULL OR sm.shift_date IS NULL)
                ORDER BY s.shift_date ASC, s.group_name ASC, u.id ASC
                """,
                (start_s, end_s, boluk_filter),
            ).fetchall()
        else:
            all_rows = db.execute(
                """
                SELECT s.shift_date, s.group_name, u.id AS user_id, u.name AS user_name
                FROM shifts s
                JOIN users u ON u.id = s.user_id
                LEFT JOIN shift_meta sm ON sm.shift_date = s.shift_date
                WHERE s.shift_date >= ? AND s.shift_date < ?
                  AND (sm.confirmed_at IS NOT NULL OR sm.shift_date IS NULL)
                ORDER BY s.shift_date ASC, s.group_name ASC, u.id ASC
                """,
                (start_s, end_s),
            ).fetchall()
        matrix_users = users
        matrix = {u["id"]: {} for u in matrix_users}
        for r in all_rows:
            matrix[r["user_id"]][r["shift_date"]] = r["group_name"]
    elif user_id:
        selected_user = db.execute("SELECT * FROM users WHERE id=?", (user_id,)).fetchone()
        if selected_user:
            history_rows = get_user_history(user_id, ym)
            matrix_users = [selected_user]
            matrix = {selected_user["id"]: {}}
            for r in history_rows:
                matrix[selected_user["id"]][r["shift_date"]] = r["group_name"]

    return render_template(
        "history.html",
        ym=ym,
        users=users,
        show_all=show_all,
        user_id=user_id,
        boluk_filter=boluk_filter,
        selected_user=selected_user,
        history_rows=history_rows,
        all_rows=all_rows,
        days=days,
        matrix_users=matrix_users,
        matrix=matrix,
    )


@app.route("/console", methods=["GET", "POST"])
def console():
    if not session.get("console_auth"):
        if request.method == "POST":
            pin = request.form.get("pin") or ""
            if pin == CONSOLE_PIN:
                session["console_auth"] = True
                return redirect(url_for("console"))
        return render_template("console_login.html")

    db = get_db()
    output = ""
    cmd_text = ""
    if request.method == "POST":
        cmd_text = request.form.get("commands") or ""
        output = execute_console_commands(cmd_text, db)

    return render_template(
        "console.html",
        output=output,
        cmd_text=cmd_text,
    )


@app.route("/console/logout")
def console_logout():
    session.pop("console_auth", None)
    return redirect(url_for("console"))


def simulate_test_schedule(ym: str, reset: bool, seed: Optional[int] = None):
    global _rng
    old_rng = _rng
    if seed is not None:
        _rng = random.Random(seed)
    start_s, end_s = month_range(ym)
    start = datetime.strptime(start_s, "%Y-%m-%d").date()
    end = datetime.strptime(end_s, "%Y-%m-%d").date()
    total_days = min(30, (end - start).days)

    days = [start + timedelta(days=i) for i in range(total_days)]

    db = get_db()
    settings = load_runtime_settings(db)
    status_days = settings["status_days"]
    min_gap_days = settings["min_gap_days"]
    yx_cooldown_choices = settings["yx_cooldown_choices"]
    boluk2_enabled = settings.get("boluk2_enabled", False)
    boluk2_cycle_days = settings.get("boluk2_cycle_days", BOLUK2_CYCLE_DAYS)
    ttm_cycle_days = settings.get("ttm_cycle_days", TTM_CYCLE_DAYS)
    weekend_leave_ids = get_weekend_leave_ids(db)
    settings["weekend_leave_ids"] = weekend_leave_ids
    users = db.execute("SELECT * FROM users ORDER BY id ASC").fetchall()
    prev_ym = previous_month(ym)
    settings["prev_month_counts"] = get_monthly_counts(prev_ym)
    settings["max_yx_per_day"] = MAX_YX_PER_DAY
    sim_users = [
        {
            "id": u["id"],
            "name": u["name"],
            "role": u["role"],
            "status": u["status"],
            "boluk": u["boluk"],
            "last_selected": u["last_selected"],
            "last_group": u["last_group"],
            "rotation_score": u["rotation_score"],
            "yx_next_eligible": u["yx_next_eligible"],
            "nb_next_eligible": u["nb_next_eligible"],
            "cycle_mask": u["cycle_mask"] if "cycle_mask" in u.keys() else 0,
            "cycle_started": u["cycle_started"] if "cycle_started" in u.keys() else None,
        }
        for u in users
    ]
    settings["user_status_by_id"] = {
        u["id"]: normalize_status_code(u["status"]) for u in sim_users
    }
    if reset:
        for u in sim_users:
            u["last_selected"] = None
            u["last_group"] = None
            u["rotation_score"] = 0
            u["yx_next_eligible"] = None
            u["nb_next_eligible"] = None
            u["cycle_mask"] = 0
            u["cycle_started"] = None
    normal_candidates = [u for u in sim_users if u["role"] == "normal"]
    eligible_ids = {
        u["id"]
        for u in normal_candidates
        if normalize_status_code(u["status"]) not in ALWAYS_EXCLUDED_STATUSES
    }
    monthly_counts = {} if reset else dict(get_monthly_counts(ym))
    monthly_group_counts = {} if reset else get_monthly_group_counts(ym)
    group_last, category_last = ({}, {}) if reset else build_user_group_history(db)
    fill_counts: Dict[int, int] = {}

    group_order = ["bas"] + [g for g, _ in GROUP_SLOTS]
    users_view = [{"id": u["id"], "name": u["name"], "boluk": u["boluk"]} for u in sim_users]
    matrix: Dict[int, Dict[str, str]] = {u["id"]: {} for u in sim_users}
    last_boluk2_cycle = get_boluk2_last_cycle(db) if not reset else None
    last_ttm_cycle = parse_date(get_setting(db, "ttm_last_cycle", None)) if not reset else None
    last_cycle_boluk1_ids: set[int] = (
        set(get_prev_cycle_boluk1_ids(db, last_boluk2_cycle)) if last_boluk2_cycle else set()
    )

    for day_index, d in enumerate(days):
        picks: Dict[str, List[int]] = {g: [] for g, _ in GROUP_SLOTS}
        picks["bas"] = []
        selected_by_group = {g: [] for g in group_order}
        cycle_day = boluk2_enabled and is_boluk2_cycle_day(
            d, last_boluk2_cycle, boluk2_cycle_days
        )
        ttm_cycle_day = is_boluk2_cycle_day(
            d, last_ttm_cycle, ttm_cycle_days
        )
        rolling_group_counts = build_sim_rolling_group_counts(
            normal_candidates, days, matrix, day_index
        )
        settings["rolling_group_counts"] = rolling_group_counts
        settings["rolling_group_avg"] = compute_rolling_group_avg(
            rolling_group_counts, eligible_ids
        )
        boluk2_active_all = [
            u
            for u in normal_candidates
            if u["boluk"] == 2 and normalize_status_code(u["status"]) == "aktiv"
        ]
        effective_cycle_day = cycle_day and bool(boluk2_active_all)

        # bas selection (role=bas only)
        bas_candidates = [
            u
            for u in sim_users
            if u["role"] == "bas"
            and normalize_status_code(u["status"]) not in ALWAYS_EXCLUDED_STATUSES
            and status_allowed_today(u["status"], d, status_days)
        ]
        if boluk2_enabled:
            preferred_boluk = 2 if effective_cycle_day else 1
            filtered = [u for u in bas_candidates if u["boluk"] == preferred_boluk]
            bas_candidates = filtered or bas_candidates
        else:
            bas_candidates = [u for u in bas_candidates if u["boluk"] == 1]
        chosen_bas = pick_bas_user(
            bas_candidates,
            d,
            status_days,
            min_gap_days,
            settings,
            weekend_leave_ids,
            monthly_counts,
        )
        if chosen_bas:
            picks["bas"].append(chosen_bas["id"])
            selected_by_group["bas"].append(chosen_bas)

        # normal selections (10 users)
        daily_candidates = normal_candidates
        if not boluk2_enabled or not effective_cycle_day:
            daily_candidates = [u for u in normal_candidates if u["boluk"] == 1]
        quota = monthly_quota(ym, daily_candidates, status_days, total_days=total_days)
        if quota:
            settings["monthly_quota"] = quota

        forced_ids: set[int] = set()
        fill_needed = 0
        fill_pool: List[dict] = []
        fill_exclude_ids: set[int] = set()
        force_only = False

        ttm_candidates = [
            u for u in daily_candidates if normalize_status_code(u["status"]) == "ttm"
        ]
        ttm_forced_ids: set[int] = set()
        ttm_block_ids: set[int] = set()
        if ttm_candidates:
            if ttm_cycle_day:
                if all(
                    day_eligible_common(
                        u, d, status_days, min_gap_days, settings, weekend_leave_ids
                    )
                    for u in ttm_candidates
                ):
                    ttm_forced_ids = {u["id"] for u in ttm_candidates}
                else:
                    ttm_block_ids = {u["id"] for u in ttm_candidates}
                    daily_candidates = [
                        u for u in daily_candidates if u["id"] not in ttm_block_ids
                    ]
            else:
                ttm_block_ids = {u["id"] for u in ttm_candidates}
                daily_candidates = [
                    u for u in daily_candidates if u["id"] not in ttm_block_ids
                ]

        if effective_cycle_day:
            boluk2_active = boluk2_active_all
            forced_ids = {u["id"] for u in boluk2_active}
        forced_ids |= ttm_forced_ids
        if len(forced_ids) > NORMAL_SLOTS and ttm_forced_ids:
            forced_ids -= ttm_forced_ids
            ttm_block_ids |= ttm_forced_ids
            daily_candidates = [u for u in daily_candidates if u["id"] not in ttm_forced_ids]
            ttm_forced_ids = set()

        if len(forced_ids) < NORMAL_SLOTS:
            fill_needed = NORMAL_SLOTS - len(forced_ids)
            if effective_cycle_day:
                fill_pool = [
                    u
                    for u in normal_candidates
                    if u["boluk"] == 1 and u["id"] not in ttm_block_ids
                ]
                fill_exclude_ids = last_cycle_boluk1_ids
            else:
                fill_pool = [u for u in daily_candidates if u["id"] not in ttm_block_ids]
        force_only = bool(forced_ids) and fill_needed == 0

        for group in GROUP_PICK_ORDER:
            chosen = None
            if forced_ids:
                pool = [u for u in daily_candidates if u["id"] in forced_ids]
                if pool:
                    try:
                        chosen = pick_weighted_random(
                            pool,
                            group,
                            d,
                            picks,
                            monthly_counts,
                            prefer_yx=True,
                            group_last=group_last,
                            category_last=category_last,
                            monthly_group_counts=monthly_group_counts,
                            settings=settings,
                            allow_cooldown_break=True,
                            allow_rotation_break=False,
                        )
                    except RuntimeError:
                        chosen = None
                    if chosen is None:
                        try:
                            chosen = pick_weighted_random(
                                pool,
                                group,
                                d,
                                picks,
                                monthly_counts,
                                prefer_yx=True,
                                group_last=group_last,
                                category_last=category_last,
                                monthly_group_counts=monthly_group_counts,
                                settings=settings,
                                allow_cooldown_break=True,
                                allow_rotation_break=True,
                            )
                        except RuntimeError:
                            chosen = None
                    if chosen is not None:
                        forced_ids.discard(chosen["id"])
            if chosen is None and fill_needed > 0:
                pool = [u for u in fill_pool if u["id"] not in fill_exclude_ids]
                if not pool:
                    pool = fill_pool
                if pool:
                    if effective_cycle_day:
                        pool = fill_pool_priority(pool, fill_counts)
                    try:
                        chosen = pick_weighted_random(
                            pool,
                            group,
                            d,
                            picks,
                            monthly_counts,
                            prefer_yx=True,
                            group_last=group_last,
                            category_last=category_last,
                            monthly_group_counts=monthly_group_counts,
                            settings=settings,
                        )
                    except RuntimeError:
                        try:
                            chosen = pick_weighted_random(
                                pool,
                                group,
                                d,
                                picks,
                                monthly_counts,
                                prefer_yx=True,
                                group_last=group_last,
                                category_last=category_last,
                                monthly_group_counts=monthly_group_counts,
                                settings=settings,
                                allow_cooldown_break=True,
                                allow_rotation_break=False,
                            )
                        except RuntimeError:
                            chosen = None
                    if chosen is None:
                        try:
                            chosen = pick_weighted_random(
                                pool,
                                group,
                                d,
                                picks,
                                monthly_counts,
                                prefer_yx=True,
                                group_last=group_last,
                                category_last=category_last,
                                monthly_group_counts=monthly_group_counts,
                                settings=settings,
                                allow_cooldown_break=True,
                                allow_rotation_break=True,
                            )
                        except RuntimeError:
                            chosen = None
                if chosen is not None:
                    fill_needed -= 1
            if chosen is None and not force_only:
                try:
                    chosen = pick_weighted_random(
                        daily_candidates,
                        group,
                        d,
                        picks,
                        monthly_counts,
                        prefer_yx=True,
                        group_last=group_last,
                        category_last=category_last,
                        monthly_group_counts=monthly_group_counts,
                        settings=settings,
                    )
                except RuntimeError:
                    try:
                        chosen = pick_weighted_random(
                            daily_candidates,
                            group,
                            d,
                            picks,
                            monthly_counts,
                            prefer_yx=True,
                            group_last=group_last,
                            category_last=category_last,
                            monthly_group_counts=monthly_group_counts,
                            settings=settings,
                            allow_cooldown_break=True,
                            allow_rotation_break=False,
                        )
                    except RuntimeError:
                        chosen = None
                if chosen is None:
                    try:
                        chosen = pick_weighted_random(
                            daily_candidates,
                            group,
                            d,
                            picks,
                            monthly_counts,
                            prefer_yx=True,
                            group_last=group_last,
                            category_last=category_last,
                            monthly_group_counts=monthly_group_counts,
                            settings=settings,
                            allow_cooldown_break=True,
                            allow_rotation_break=True,
                        )
                    except RuntimeError:
                        chosen = None
            if chosen is None:
                continue
            if forced_ids and chosen["id"] in forced_ids:
                forced_ids.discard(chosen["id"])
            picks[group].append(chosen["id"])
            selected_by_group[group].append(chosen)

        if effective_cycle_day or ttm_forced_ids:
            users_by_id = {u["id"]: u for u in sim_users}
            selected_ids = {uid for ids in picks.values() for uid in ids}
            forced_expected: set[int] = set(ttm_forced_ids)
            if effective_cycle_day:
                forced_expected |= {u["id"] for u in boluk2_active}
            missing_forced = forced_expected - selected_ids
            if missing_forced:
                swap_in_forced_users(picks, missing_forced, users_by_id, monthly_group_counts)
            selected_by_group = {
                g: [users_by_id[uid] for uid in picks.get(g, [])]
                for g in group_order
            }

        ds = d.strftime("%Y-%m-%d")
        for group in group_order:
            for u in selected_by_group[group]:
                matrix[u["id"]][ds] = group

        if effective_cycle_day:
            boluk2_active_ids = {u["id"] for u in boluk2_active}
            selected_ids = {u["id"] for group in group_order for u in selected_by_group[group] if group != "bas"}
            if boluk2_active_ids and boluk2_active_ids.issubset(selected_ids):
                last_boluk2_cycle = d
                last_cycle_boluk1_ids = {
                    u["id"]
                    for group in group_order
                    for u in selected_by_group[group]
                    if group != "bas" and u["boluk"] == 1
                }
        if ttm_forced_ids:
            selected_ids = {
                u["id"]
                for group in group_order
                for u in selected_by_group[group]
                if group != "bas"
            }
            if ttm_forced_ids.issubset(selected_ids):
                last_ttm_cycle = d

        for group in group_order:
            for u in selected_by_group[group]:
                u["last_selected"] = ds
                u["last_group"] = group
                u["rotation_score"] = float(u["rotation_score"] or 0) + 1
                monthly_counts[u["id"]] = monthly_counts.get(u["id"], 0) + 1
                monthly_group_counts.setdefault(u["id"], {})
                monthly_group_counts[u["id"]][group] = (
                    monthly_group_counts[u["id"]].get(group, 0) + 1
                )
                prev_mask = user_cycle_mask(u)
                new_mask = next_cycle_mask(prev_mask, group)
                if prev_mask == 0 and new_mask != 0 and not u.get("cycle_started"):
                    u["cycle_started"] = ds
                if prev_mask != 0 and new_mask == 0:
                    u["cycle_started"] = ds
                u["cycle_mask"] = new_mask
                if normalize_status_code(u["status"]) == "yx":
                    cooldown = _rng.choice(yx_cooldown_choices)
                    u["yx_next_eligible"] = (d + timedelta(days=cooldown)).strftime(
                        "%Y-%m-%d"
                    )
                if group in {"nbm", "nbm2"}:
                    cooldown = _rng.choice(settings.get("nb_cooldown_choices", NB_COOLDOWN_CHOICES))
                    u["nb_next_eligible"] = (d + timedelta(days=cooldown)).strftime(
                        "%Y-%m-%d"
                    )
                group_last.setdefault(u["id"], {})[group] = d
                cat = group_category(group)
                category_last.setdefault(u["id"], {})[cat] = d

        if effective_cycle_day:
            for group in group_order:
                if group == "bas":
                    continue
                for u in selected_by_group[group]:
                    if user_boluk(u) == 1 and u.get("role") == "normal":
                        fill_counts[u["id"]] = fill_counts.get(u["id"], 0) + 1

        # Simulate daily decay for long-term fairness
        for u in normal_candidates:
            rs = float(u["rotation_score"] or 0)
            u["rotation_score"] = max(0.0, rs - 0.2)

    _rng = old_rng
    return days, users_view, matrix, total_days


def build_daily_groups(days, users_view, matrix):
    order = ["bas", "g1", "g2", "g3", "nbm", "nbm2", "p1", "p2", "p3"]
    daily = []
    for d in days:
        ds = d.strftime("%Y-%m-%d")
        groups = {g: [] for g in order}
        for u in users_view:
            g = matrix.get(u["id"], {}).get(ds)
            if g:
                groups.setdefault(g, []).append(u["name"])
        daily.append((ds, groups))
    return daily


def get_test_saved_at(db: sqlite3.Connection, ym: str) -> Optional[str]:
    row = db.execute(
        "SELECT saved_at FROM test_meta WHERE month=?", (ym,)
    ).fetchone()
    return row["saved_at"] if row else None


def load_saved_test_schedule(ym: str):
    db = get_db()
    saved_at = get_test_saved_at(db, ym)
    if not saved_at:
        return None

    start_s, end_s = month_range(ym)
    start = datetime.strptime(start_s, "%Y-%m-%d").date()
    end = datetime.strptime(end_s, "%Y-%m-%d").date()
    total_days = min(30, (end - start).days)
    days = [start + timedelta(days=i) for i in range(total_days)]

    users = db.execute("SELECT id, name, boluk FROM users ORDER BY id ASC").fetchall()
    users_view = [{"id": u["id"], "name": u["name"], "boluk": u["boluk"]} for u in users]
    matrix: Dict[int, Dict[str, str]] = {u["id"]: {} for u in users_view}

    rows = db.execute(
        """
        SELECT shift_date, group_name, user_id
        FROM test_shifts
        WHERE month=?
        """,
        (ym,),
    ).fetchall()
    for r in rows:
        if r["user_id"] in matrix:
            matrix[r["user_id"]][r["shift_date"]] = r["group_name"]

    return days, users_view, matrix, total_days, saved_at


def save_test_schedule(
    db: sqlite3.Connection,
    ym: str,
    days,
    users_view,
    matrix,
) -> str:
    db.execute("DELETE FROM test_shifts WHERE month=?", (ym,))
    for u in users_view:
        uid = u["id"]
        u_map = matrix.get(uid, {})
        for d in days:
            ds = d.strftime("%Y-%m-%d")
            g = u_map.get(ds)
            if not g:
                continue
            db.execute(
                """
                INSERT INTO test_shifts (month, shift_date, group_name, user_id)
                VALUES (?, ?, ?, ?)
                """,
                (ym, ds, g, uid),
            )

    saved_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    db.execute(
        """
        INSERT INTO test_meta (month, saved_at)
        VALUES (?, ?)
        ON CONFLICT(month) DO UPDATE SET saved_at=excluded.saved_at
        """,
        (ym, saved_at),
    )
    db.commit()
    return saved_at


def clear_test_schedule(db: sqlite3.Connection, ym: str) -> None:
    db.execute("DELETE FROM test_shifts WHERE month=?", (ym,))
    db.execute("DELETE FROM test_meta WHERE month=?", (ym,))
    db.commit()


@app.route("/test")
def test_view():
    ym = request.args.get("month") or date.today().strftime("%Y-%m")
    reset = request.args.get("reset") == "1"
    seed = request.args.get("seed", type=int)
    if reset and seed is None:
        seed = _rng.randint(1, 9999999)
        return redirect(url_for("test_view", month=ym, reset=1, seed=seed))
    db = get_db()
    test_saved_at = None
    saved = None
    if not reset:
        saved = load_saved_test_schedule(ym)
    if saved:
        days, users_view, matrix, total_days, test_saved_at = saved
    else:
        days, users_view, matrix, total_days = simulate_test_schedule(ym, reset, seed)
        if not reset:
            test_saved_at = get_test_saved_at(db, ym)
    seed_val = "" if seed is None else seed
    boluk1_users = [u for u in users_view if u.get("boluk") == 1]
    boluk2_users = [u for u in users_view if u.get("boluk") == 2]
    return render_template(
        "test.html",
        ym=ym,
        days=days,
        users=users_view,
        boluk1_users=boluk1_users,
        boluk2_users=boluk2_users,
        matrix=matrix,
        total_days=total_days,
        reset=reset,
        seed=seed_val,
        test_saved_at=test_saved_at,
    )


@app.route("/test/save", methods=["POST"])
def test_save():
    ym = request.form.get("month") or date.today().strftime("%Y-%m")
    reset = request.form.get("reset") == "1"
    seed_raw = (request.form.get("seed") or "").strip()
    seed = int(seed_raw) if seed_raw.isdigit() else None

    days, users_view, matrix, _total = simulate_test_schedule(ym, reset, seed)
    db = get_db()
    save_test_schedule(db, ym, days, users_view, matrix)
    flash("Test yadda saxlandi.", "ok")
    return redirect(url_for("test_view", month=ym))


@app.route("/test/clear", methods=["POST"])
def test_clear():
    ym = request.form.get("month") or date.today().strftime("%Y-%m")
    db = get_db()
    clear_test_schedule(db, ym)
    flash("Test temizlendi.", "ok")
    return redirect(url_for("test_view", month=ym))


def add_group_docx(doc: "Document", label: str, names: List[str]) -> None:
    if not names:
        doc.add_paragraph(f"{label}: -")
        return
    text = f"{label}: " + "\n".join(names)
    doc.add_paragraph(text)


@app.route("/export_history.docx")
def export_history_docx():
    if Document is None:
        return "python-docx not installed. Run: pip install python-docx", 400

    ym = request.args.get("month") or date.today().strftime("%Y-%m")
    user_id_raw = request.args.get("user_id")
    show_all = not user_id_raw or user_id_raw == "all"
    db = get_db()

    doc = Document()
    if show_all:
        start_s, end_s = month_range(ym)
        start = datetime.strptime(start_s, "%Y-%m-%d").date()
        end = datetime.strptime(end_s, "%Y-%m-%d").date()
        total_days = (end - start).days
        days = [start + timedelta(days=i) for i in range(total_days)]

        rows = db.execute(
            """
            SELECT s.shift_date, s.group_name, u.name AS user_name
            FROM shifts s
            JOIN users u ON u.id = s.user_id
            LEFT JOIN shift_meta sm ON sm.shift_date = s.shift_date
            WHERE s.shift_date >= ? AND s.shift_date < ?
              AND (sm.confirmed_at IS NOT NULL OR sm.shift_date IS NULL)
            ORDER BY s.shift_date ASC, s.group_name ASC, u.id ASC
            """,
            (start_s, end_s),
        ).fetchall()

        by_day: Dict[str, Dict[str, List[str]]] = {}
        for r in rows:
            by_day.setdefault(r["shift_date"], {})
            by_day[r["shift_date"]].setdefault(r["group_name"], []).append(r["user_name"])

        doc.add_heading(f"Naryad History - {ym}", level=1)
        order = ["bas", "g1", "g2", "g3", "nbm", "nbm2", "p1", "p2", "p3"]
        for d in days:
            ds = d.strftime("%Y-%m-%d")
            doc.add_heading(d.strftime("%d.%m.%Y"), level=2)
            groups = by_day.get(ds, {})
            for g in order:
                names = groups.get(g, [])
                label = GROUP_LABELS.get(g, g)
                add_group_docx(doc, label, names)
            doc.add_paragraph("")
    else:
        try:
            user_id = int(user_id_raw)
        except Exception:
            return "Invalid user_id.", 400

        user = db.execute("SELECT name FROM users WHERE id=?", (user_id,)).fetchone()
        if not user:
            return "User not found.", 404

        rows = get_user_history(user_id, ym)
        doc.add_heading(f"{user['name']} - {ym} History", level=1)
        if not rows:
            doc.add_paragraph("No entries.")
        for r in rows:
            ds = r["shift_date"]
            try:
                d = datetime.strptime(ds, "%Y-%m-%d").date()
                d_label = d.strftime("%d.%m.%Y")
            except Exception:
                d_label = ds
            glabel = GROUP_LABELS.get(r["group_name"], r["group_name"])
            doc.add_paragraph(f"{d_label} - {glabel}")

    suffix = "all" if show_all else "user"
    out_path = os.path.join(EXPORT_DIR, f"history_{ym.replace('-', '_')}_{suffix}.docx")
    doc.save(out_path)
    return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))


@app.route("/export_test.docx")
def export_test_docx():
    if Document is None:
        return "python-docx not installed. Run: pip install python-docx", 400

    ym = request.args.get("month") or date.today().strftime("%Y-%m")
    reset = request.args.get("reset") == "1"
    seed = request.args.get("seed", type=int)
    saved = None
    if not reset:
        saved = load_saved_test_schedule(ym)
    if saved:
        days, users_view, matrix, _total, _saved_at = saved
    else:
        days, users_view, matrix, _total = simulate_test_schedule(ym, reset, seed)
    daily = build_daily_groups(days, users_view, matrix)

    doc = Document()
    doc.add_heading(f"Test Simulation - {ym}", level=1)
    order = ["bas", "g1", "g2", "g3", "nbm", "nbm2", "p1", "p2", "p3"]
    for ds, groups in daily:
        try:
            d = datetime.strptime(ds, "%Y-%m-%d").date()
            d_label = d.strftime("%d.%m.%Y")
        except Exception:
            d_label = ds
        doc.add_heading(d_label, level=2)
        for g in order:
            names = groups.get(g, [])
            label = GROUP_LABELS.get(g, g)
            add_group_docx(doc, label, names)
        doc.add_paragraph("")

    suffix = "reset" if reset else "seeded"
    out_path = os.path.join(EXPORT_DIR, f"test_{ym.replace('-', '_')}_{suffix}.docx")
    doc.save(out_path)
    return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))


@app.route("/export.docx")
def export_docx():
    if Document is None:
        return "python-docx quraşdırılmayıb. Terminalda: pip install python-docx", 400

    d = request.args.get("d") or date.today().strftime("%Y-%m-%d")
    shift_date = datetime.strptime(d, "%Y-%m-%d").date()
    shift_next_date = shift_date + timedelta(days=1)
    data = load_shift(shift_date)

    doc = Document()
    doc.add_heading(
        f"I Mühafizə Bölüyü Naryad Cədvəli - {shift_date.strftime("%d.%m.%Y")}", level=1
    )
    doc.add_paragraph(
        f"Vaxtı: {SHIFT_START_HOUR:02d}:00 -> növbəti gün {SHIFT_START_HOUR:02d}:00 ({shift_date.strftime("%d.%m.%Y")} -> {shift_next_date.strftime("%d.%m.%Y")})"
    )

    bas = data.get("bas", [])
    if bas:
        doc.add_paragraph(f"Tabor Növbətçisi: {bas[0]["name"]}")
    else:
        doc.add_paragraph("Tabor Növbətçisi: (təyin olunmayıb)")

    doc.add_paragraph("")

    order = ["g1", "g2", "g3", "nbm", "nbm2", "p1", "p2", "p3"]
    for gname in order:
        glabel = GROUP_LABELS.get(gname, gname)
        people = data.get(gname, [])
        names = [p["name"] for p in people]
        add_group_docx(doc, glabel, names)

    out_path = os.path.join(EXPORT_DIR, f"naryad_{shift_date.strftime('%Y_%m_%d')}.docx")
    doc.save(out_path)
    return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))


if __name__ == "__main__":
    init_db()
    app.run(host="127.0.0.1", port=8888, debug=False)



